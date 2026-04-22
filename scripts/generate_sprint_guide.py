"""
Generates docs/Sprint-Tasks-Implementation-Guide.docx.

Run:  python scripts/generate_sprint_guide.py

Requires: python-docx (pip install python-docx).
The script only writes the .docx. It does not touch any Assets/ files.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Output path
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "Sprint-Tasks-Implementation-Guide.docx"

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
CODE_BG = "F2F2F2"
WARN_BG = "FFF4E5"
NOTE_BG = "E8F4FD"
TIP_BG = "EAF5EA"


def _shade(element, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    element.append(shd)


def add_title_block(doc, title, subtitle, meta_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = BODY_FONT
        run.font.size = Pt(11)

    doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    return p


def add_callout(doc, kind, text):
    colors = {"WARNING": WARN_BG, "NOTE": NOTE_BG, "TIP": TIP_BG}
    bg = colors.get(kind, NOTE_BG)
    p = doc.add_paragraph()
    run = p.add_run(f"{kind}: ")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.name = BODY_FONT
    run2.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    _shade(pPr, bg)


def add_file_path(doc, path):
    p = doc.add_paragraph()
    run = p.add_run(f"File: {path}")
    run.italic = True
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_code(doc, code):
    lines = code.replace("\t", "    ").split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
        pPr = p._p.get_or_add_pPr()
        _shade(pPr, CODE_BG)
    doc.add_paragraph()


def add_commit(doc, msg):
    p = doc.add_paragraph()
    run = p.add_run("COMMIT: ")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run2 = p.add_run(msg)
    run2.italic = True
    run2.font.name = BODY_FONT
    run2.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = BODY_FONT
            run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = w
    doc.add_paragraph()


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = BODY_FONT
        run.font.size = Pt(11)


def add_checklist(doc, items):
    for it in items:
        p = doc.add_paragraph()
        r = p.add_run("\u2610  ")
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r2 = p.add_run(it)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(11)


def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ------------------------------------------------------------------ front
    add_title_block(
        doc,
        "SALINLAHI",
        "Sprint Tasks Implementation Guide",
        [
            "Four self-contained tasks. Pick any order.",
            "Owner: chdean-09",
            "April 2026",
        ],
    )

    add_h2(doc, "How to use this guide")
    add_para(
        doc,
        "Each task is a complete unit of work tracked by a single Jira ticket. Every refactor, "
        "new file, and scene edit that the task needs lives inside that task's section. Pick "
        "the task you want to work on, follow its sub-steps top-to-bottom, and commit after "
        "each sub-step with the exact commit message shown. There is no shared foundation "
        "phase.",
    )
    add_callout(
        doc,
        "NOTE",
        "Every commit message includes the SALIN ticket it belongs to (SALIN-40, SALIN-54, "
        "SALIN-68, or SALIN-69). Follow the project's Jira convention: "
        "type(scope): SALIN-XX description.",
    )
    add_callout(
        doc,
        "WARNING",
        "Two tasks touch ActiveEnemyTracker.cs for different reasons (SALIN-40 refactors "
        "FindAllWithCharacter; SALIN-68 adds ActiveCount / IsClear). And two tasks touch "
        "Enemy.cs (SALIN-54 fixes double-Unregister and adds the speed-buff API; SALIN-68 "
        "reads IsBoss). Whichever task lands second just rebases cleanly on the first — no "
        "shared code path conflicts, but expect a trivial merge if both branches are live.",
    )

    add_h2(doc, "Master Step Table")
    add_table(
        doc,
        ["Ticket", "What It Ships", "Depends On"],
        [
            (
                "SALIN-40  AOE Burst Mechanic",
                "EventBus.OnAOETriggered + ActiveEnemyTracker allocation-free query refactor + "
                "AOEResolver MonoBehaviour that mass-defeats 3+ same-character enemies on screen "
                "at draw time. Never targets bosses.",
                "Nothing.",
            ),
            (
                "SALIN-54  Pensionado + General",
                "Enemy.Defeat single-Unregister fix + EnemyDataSO variant fields + Enemy "
                "speed-buff API + PensionadoMover + GeneralAura + prefabs + data assets.",
                "Nothing.",
            ),
            (
                "SALIN-68  Boss Encounter System",
                "Singleton scene-local fix + EventBus boss events + GameManager.CurrentBoss + "
                "ActiveEnemyTracker.ActiveCount / IsClear + BaybayinCharacterSO.icon + "
                "WaveSpawner.SpawnWave + CombatResolver boss route + BossConfigSO + "
                "BossController + abilities + label icon row + El Inquisidor + stubs + tests.",
                "Nothing.",
            ),
            (
                "SALIN-69  Tracing Dojo",
                "RecognitionLogger.LoggingEnabled flag + TracingDojo.unity scene + character "
                "list UI + ghost stroke overlay + dojo controller + Main Menu Practice button.",
                "Nothing.",
            ),
        ],
        col_widths=[Cm(4.5), Cm(10.5), Cm(2)],
    )

    page_break(doc)

    # ------------------------------------------------------------------ tasks
    task_salin_40(doc)

    page_break(doc)
    task_salin_54(doc)

    page_break(doc)
    task_salin_68(doc)

    page_break(doc)
    task_salin_69(doc)

    # ------------------------------------------------------------------ back
    page_break(doc)
    back_matter(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote: {OUTPUT}")


# ===========================================================================
# SALIN-40  —  AOE Burst Mechanic
# ===========================================================================


def task_salin_40(doc):
    add_h1(doc, "SALIN-40  Implement AOE Burst Mechanic (3+ Same Character Mass Defeat)")
    add_para(
        doc,
        "Risk-reward pressure valve. Players can intentionally let same-character enemies "
        "accumulate on screen \u2014 they get closer to the shrine, which is dangerous \u2014 in exchange "
        "for a single-draw mass clear when three or more are on screen together. The mechanic "
        "is purely spatial and per-draw: no streaks, no memory between draws.",
    )
    add_para(doc, "This ticket ships three things: the missing event, a non-allocating tracker query, and the resolver itself.")

    # ---------- 1.1 EventBus ----------
    add_h2(doc, "1.1  Add OnAOETriggered event to EventBus")
    add_para(
        doc,
        "Spec'd in docs/system/03_Core_Systems.md but not declared in code today. "
        "AOEResolver raises it after damage lands, so subscribers (HUD badges, analytics, "
        "SFX) can react.",
    )
    add_file_path(doc, "Assets/Scripts/Core/EventBus.cs")
    add_para(doc, "Add the event and Raise helper alongside the existing declarations:")
    add_code(
        doc,
        """// Raised once per AOE resolution, after damage is applied to every matching enemy.
public static event Action<List<BaybayinCharacterSO>> OnAOETriggered;

public static void RaiseAOETriggered(List<BaybayinCharacterSO> defeated)
{
    OnAOETriggered?.Invoke(defeated);
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "Add `using System.Collections.Generic;` to the top of the file if it is not already "
        "imported \u2014 the generic List<> type needs it. The compiler will flag this if you miss it.",
    )
    add_commit(doc, "feat(events): SALIN-40 add OnAOETriggered event to EventBus")

    # ---------- 1.2 ActiveEnemyTracker refactor ----------
    add_h2(doc, "1.2  ActiveEnemyTracker refactor \u2014 already done in SALIN-89")
    add_callout(
        doc,
        "NOTE",
        "This step is ALREADY COMPLETE on main. SALIN-89 shipped the allocation-free "
        "FindAllWithCharacter (reuses _characterMatchBuffer), deduped the prune logic into "
        "CleanupStaleEntries(), and added ActiveCount / IsClear / GetActiveEnemiesSnapshot(). "
        "Skip the code edit \u2014 just verify and move on.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Enemy/ActiveEnemyTracker.cs")
    add_para(doc, "Confirm the class currently looks like this (do not edit):")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;

public class ActiveEnemyTracker : Singleton<ActiveEnemyTracker>
{
    private readonly List<Enemy> _activeEnemies = new List<Enemy>();
    private readonly List<Enemy> _characterMatchBuffer = new List<Enemy>();

    public int ActiveCount { get { CleanupStaleEntries(); return _activeEnemies.Count; } }
    public bool IsClear => ActiveCount == 0;

    public List<Enemy> GetActiveEnemiesSnapshot() { ... }
    public Enemy FindClosestToBase(string characterID) { ... }
    public List<Enemy> FindAllWithCharacter(string characterID) { ... reuses _characterMatchBuffer ... }
    private void CleanupStaleEntries() { ... }
}""",
    )
    add_callout(
        doc,
        "WARNING",
        "FindAllWithCharacter returns the shared _characterMatchBuffer. Do NOT cache the result "
        "across frames or across subsequent calls \u2014 the next call overwrites it. AOEResolver "
        "reads it immediately inside the OnCharacterRecognized handler and never stores it, "
        "which is the correct usage. The return type is the concrete List<Enemy>, not "
        "IReadOnlyList<Enemy> \u2014 do not treat it as snapshot-safe.",
    )
    add_para(
        doc,
        "AC-10 and AC-11 are already satisfied by SALIN-89. No commit for step 1.2. "
        "Skip ahead to 1.3.",
    )


    # ---------- 1.3 AOEResolver ----------
    add_h2(doc, "1.3  Create AOEResolver")
    add_para(
        doc,
        "A plain MonoBehaviour (not a Singleton) living in the Gameplay scene. Subscribes to "
        "OnCharacterRecognized in OnEnable, unsubscribes in OnDisable. On each recognition: "
        "query matching enemies, count them, mass-defeat all non-bosses if >= 3.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Combat/AOEResolver.cs")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;

public class AOEResolver : MonoBehaviour
{
    private const int AOE_THRESHOLD = 3;

    [SerializeField] private CharacterRegistrySO _registry;
    [SerializeField] private GameObject _aoeFlashVfxPrefab; // optional

    private readonly List<BaybayinCharacterSO> _defeatedBuffer
        = new List<BaybayinCharacterSO>(16);

    private void OnEnable()
    {
        EventBus.OnCharacterRecognized += OnCharacterRecognized;
    }

    private void OnDisable()
    {
        EventBus.OnCharacterRecognized -= OnCharacterRecognized;
    }

    private void OnCharacterRecognized(string characterID)
    {
        // AC-7: if a boss is consuming this draw, the CombatResolver short-circuit
        // (added by SALIN-68) has already returned before us. Defensive: if a boss is
        // targetable and this char is required by the current phase, skip AOE.
        // Safe no-op when no boss feature is installed yet \u2014 GameManager.CurrentBoss
        // will simply be null.
        var boss = GameManager.Instance != null
            ? GameManager.Instance.CurrentBoss
            : null;
        if (boss != null && boss.IsTargetable && BossWouldAccept(boss, characterID))
            return;

        var matches = ActiveEnemyTracker.Instance.FindAllWithCharacter(characterID);
        if (matches == null || matches.Count < AOE_THRESHOLD) return; // AC-4

        _defeatedBuffer.Clear();
        var charSO = _registry.GetByID(characterID);

        // Iterate on a snapshot \u2014 TakeDamage triggers Defeat \u2192 ReturnToPool \u2192
        // Unregister, which mutates the tracker's list under us.
        var snapshot = new List<Enemy>(matches);
        for (int i = 0; i < snapshot.Count; i++)
        {
            var e = snapshot[i];
            if (e == null) continue;
            if (e.IsBoss) continue; // AC-7 belt-and-suspenders
            e.TakeDamage(e.MaxHealth);
            _defeatedBuffer.Add(charSO);
        }

        if (_aoeFlashVfxPrefab != null)
            Instantiate(_aoeFlashVfxPrefab, transform.position, Quaternion.identity);

        // AC-5: raise exactly once, after damage has been applied.
        EventBus.RaiseAOETriggered(_defeatedBuffer);
    }

    private static bool BossWouldAccept(BossController boss, string characterID)
    {
        var phase = boss.CurrentPhase;
        if (phase == null || phase.requiredCharacters == null) return false;
        for (int i = 0; i < phase.requiredCharacters.Count; i++)
        {
            var c = phase.requiredCharacters[i];
            if (c != null && c.characterID == characterID) return true;
        }
        return false;
    }
}""",
    )
    add_callout(
        doc,
        "WARNING",
        "The snapshot copy (`new List<Enemy>(matches)`) IS needed. Without it, TakeDamage "
        "\u2192 ReturnToPool \u2192 Unregister mutates the tracker list mid-iteration and every other "
        "enemy gets skipped. The snapshot is the only per-AOE allocation, and it only fires "
        "when count >= 3, so it is off the hot path.",
    )
    add_callout(
        doc,
        "NOTE",
        "AOEResolver references Enemy.IsBoss and BossController. If you implement this ticket "
        "before SALIN-68 lands, stub both: add `public bool IsBoss => false;` on Enemy and an "
        "empty `public class BossController : MonoBehaviour { public bool IsTargetable => false; "
        "public BossPhase CurrentPhase => null; }` so the project compiles. SALIN-68 will "
        "replace both with the real implementations. Enemy.MaxHealth can be a shim "
        "`public int MaxHealth => Data.maxHealth;` if the field does not already exist.",
    )
    add_commit(doc, "feat(combat): SALIN-40 add AOEResolver for 3+ same-character mass defeat")

    # ---------- 1.4 Place in scene ----------
    add_h2(doc, "1.4  Place AOEResolver in the Gameplay scene")
    add_para(
        doc,
        "Open Assets/Scenes/Gameplay.unity. Create an empty GameObject at the root named "
        "'AOEResolver'. Add the AOEResolver component. Assign the CharacterRegistrySO asset "
        "to the _registry field. Leave _aoeFlashVfxPrefab empty for now \u2014 the mechanic works "
        "without VFX.",
    )
    add_callout(
        doc,
        "WARNING",
        "Do NOT add AOEResolver to the Bootstrap scene or mark it DontDestroyOnLoad. It is "
        "scene-local by design; spawning a second one when re-entering Gameplay would double-"
        "subscribe to OnCharacterRecognized and every AOE would fire twice.",
    )
    add_commit(doc, "chore(scene): SALIN-40 add AOEResolver GameObject to Gameplay scene")

    # ---------- 1.5 Optional VFX / HUD hook ----------
    add_h2(doc, "1.5  Optional: AOE flash VFX and HUD hook")
    add_para(
        doc,
        "Non-blocking polish. The mechanic is functional without either, but these sharpen "
        "the feel. Skip to verification if you want to ship bare-bones first.",
    )
    add_h3(doc, "Flash VFX")
    add_para(
        doc,
        "If you have a full-screen flash / particle burst prefab, assign it to the resolver's "
        "_aoeFlashVfxPrefab slot in the inspector. AC-8 says the VFX must not obscure the boss "
        "label icon row (added by SALIN-68). Put the VFX on a Canvas with a lower sort order "
        "than the boss HUD Canvas. For example: Gameplay FX Canvas (sort order 5) < Boss HUD "
        "Canvas (sort order 10).",
    )
    add_h3(doc, "HUD badge")
    add_para(
        doc,
        "Note: SALIN-89 split the HUD into per-widget components in "
        "Assets/Scripts/UI/HUD/ (ComboDisplay, DrawingFeedback, FocusModeIndicator, "
        "HeartDisplay, WaveDisplay). Do NOT drop the AOE listener into a monolithic "
        "HUD.cs \u2014 add a new sibling widget instead.",
    )
    add_file_path(doc, "Assets/Scripts/UI/HUD/MassClearBadge.cs")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;

public class MassClearBadge : MonoBehaviour
{
    [SerializeField] private GameObject _badgeRoot;
    [SerializeField] private TMPro.TMP_Text _label;

    private void OnEnable()  { EventBus.OnAOETriggered += OnAOE; }
    private void OnDisable() { EventBus.OnAOETriggered -= OnAOE; }

    private void OnAOE(List<BaybayinCharacterSO> defeated)
    {
        if (defeated == null || defeated.Count == 0) return;
        _badgeRoot.SetActive(true);
        _label.text = $"MASS CLEAR x{defeated.Count}";
        CancelInvoke(nameof(Hide));
        Invoke(nameof(Hide), 1.0f);
    }

    private void Hide() { _badgeRoot.SetActive(false); }
}""",
    )
    add_para(
        doc,
        "Place the MassClearBadge GameObject as a child of the HUD Canvas alongside "
        "ComboDisplay / HeartDisplay. Follow the same prefab/activation pattern SALIN-89 "
        "established for the other HUD widgets.",
    )
    add_commit(doc, "feat(ui): SALIN-40 show mass-clear HUD badge on OnAOETriggered")

    # ---------- 1.6 Verification ----------
    add_h2(doc, "1.6  Verification")
    add_table(
        doc,
        ["Test", "Expected"],
        [
            (
                "Spawn 4 Soldados all assigned BA, draw BA",
                "All 4 defeated in a single frame. OnAOETriggered raised once with 4 entries.",
            ),
            (
                "Spawn 2 Soldados assigned BA, draw BA",
                "Count < 3. Only one defeated (closest-to-base via CombatResolver). OnAOETriggered NOT raised.",
            ),
            (
                "Spawn 3 BA + 2 KA, draw BA",
                "3 BA defeated. 2 KA untouched. AC-6 satisfied.",
            ),
            (
                "Spawn 3 BA + 1 boss phase requiring BA is active, draw BA",
                "Boss routing (from SALIN-68) consumes the draw first. AOE does not fire. 3 BA still on screen.",
            ),
            (
                "Boss in intermission (untargetable), 3 BA adds on screen, draw BA",
                "AOE fires (boss is not targetable, CombatResolver does not short-circuit). 3 adds defeated.",
            ),
            (
                "Run 60s with 100 draws and background wave traffic",
                "No GC spike visible in Profiler from per-draw AOE query (benefit of the alloc-free refactor).",
            ),
            (
                "Defeat enemies via AOE, check ActiveEnemyTracker",
                "All defeated enemies unregistered and returned to pool.",
            ),
        ],
        col_widths=[Cm(8), Cm(9)],
    )
    add_commit(doc, "test(combat): SALIN-40 verify AOE threshold, exclusions, and boss priority")


# ===========================================================================
# SALIN-54  —  Pensionado + General Enemy Variants
# ===========================================================================


def task_salin_54(doc):
    add_h1(doc, "SALIN-54  Implement Pensionado (Zigzagger) and General (Commander) Enemy Variants")
    add_para(
        doc,
        "Two American-era (Chapter 2) enemies. Data-driven via EnemyDataSO \u2014 no Enemy "
        "subclasses. Pensionado descends with a sine-wave horizontal offset. General moves "
        "at 0.7\u00d7 base speed and, while alive, applies a 1.3\u00d7 speed buff to nearby "
        "American-era enemies. On General defeat, the buff is removed immediately.",
    )

    # ---------- 2.1 Enemy.Defeat fix ----------
    add_h2(doc, "2.1  Fix Enemy.Defeat double-Unregister")
    add_para(
        doc,
        "Enemy.Defeat() currently calls ActiveEnemyTracker.Unregister(this) directly, then "
        "calls ReturnToPool(), which also calls Unregister. GeneralAura (added later in this "
        "task) iterates the tracker's ActiveEnemies list; a double-remove is harmless on a "
        "List<T> today but masks real bugs. Remove the redundant call so ReturnToPool is the "
        "single source of truth.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Enemy/Enemy.cs")
    add_para(doc, "Find Defeat() and delete the direct Unregister call. The final shape should be:")
    add_code(
        doc,
        """public void Defeat()
{
    // Existing VFX, SFX, score bump stays.
    EventBus.RaiseEnemyDefeated(this);
    ReturnToPool();
    // ReturnToPool() already unregisters from ActiveEnemyTracker.
    // Do not call Unregister here \u2014 that would double-remove.
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "Any future code path that ends an enemy's life should call ReturnToPool() and "
        "nothing else. If you grep and find other direct Unregister calls in life-ending "
        "paths, fix them the same way.",
    )
    add_commit(doc, "fix(enemy): SALIN-54 remove redundant ActiveEnemyTracker.Unregister in Enemy.Defeat")

    # ---------- 2.2 Enemy speed-buff API + IsBoss ----------
    add_h2(doc, "2.2  Add speed-buff API and IsBoss helper to Enemy")
    add_para(
        doc,
        "GeneralAura applies a speed multiplier to nearby enemies and needs to remove it on "
        "defeat. A stack keyed by the source component lets multiple auras coexist without "
        "clobbering each other. Also add IsBoss so the aura and AOE can skip bosses uniformly.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Enemy/Enemy.cs")
    add_para(doc, "Add these members to the Enemy class:")
    add_code(
        doc,
        """private readonly Dictionary<object, float> _speedBuffs
    = new Dictionary<object, float>();

public bool IsBoss => Data != null && Data.isBoss;

public int MaxHealth => Data != null ? Data.maxHealth : 0;

public float EffectiveSpeed
{
    get
    {
        if (Data == null) return 0f;
        float speed = Data.moveSpeed * Data.baseSpeedMultiplier;
        foreach (var kv in _speedBuffs) speed *= kv.Value;
        return speed;
    }
}

public void ApplySpeedBuff(object source, float multiplier)
{
    _speedBuffs[source] = multiplier;
}

public void ClearSpeedBuff(object source)
{
    _speedBuffs.Remove(source);
}""",
    )
    add_para(
        doc,
        "Then grep Enemy.cs for `moveSpeed` and swap any motion-path usage to EffectiveSpeed. "
        "If the descent speed is read inside Update, change `Data.moveSpeed` to "
        "`EffectiveSpeed`:",
    )
    add_code(
        doc,
        """// Before:
transform.position += Vector3.down * (Data.moveSpeed * Time.deltaTime);

// After:
transform.position += Vector3.down * (EffectiveSpeed * Time.deltaTime);""",
    )
    add_callout(
        doc,
        "WARNING",
        "If you skip the moveSpeed \u2192 EffectiveSpeed swap, the aura will compile cleanly but "
        "have zero visible effect. Verify in-scene (step 2.9) by spawning a General near a "
        "Soldado and watching for the speed-up.",
    )
    add_callout(
        doc,
        "NOTE",
        "EnemyDataSO needs `isBoss : bool` for the IsBoss helper to compile. This is added in "
        "step 2.3 along with the other variant fields. If Enemy.cs does not already "
        "import System.Collections.Generic, add it for Dictionary<>.",
    )
    add_commit(doc, "feat(enemy): SALIN-54 add stackable speed buff API, EffectiveSpeed, and IsBoss helper to Enemy")

    # ---------- 2.3 EnemyDataSO ----------
    add_h2(doc, "2.3  Extend EnemyDataSO with variant fields")
    add_file_path(doc, "Assets/Scripts/Data/EnemyDataSO.cs")
    add_para(doc, "Add the following fields to EnemyDataSO (leave existing fields alone):")
    add_code(
        doc,
        """public enum Era { Spanish, American, Japanese }

[Header("Variant Metadata")]
public Era era = Era.Spanish;
public bool isBoss = false;

[Header("Zigzag Mover (Pensionado)")]
[Tooltip("Horizontal sine amplitude in world units. 0 disables zigzag.")]
public float zigzagAmplitude = 0f;
[Tooltip("Sine frequency in Hz.")]
public float zigzagFrequency = 0f;

[Header("Base Speed Modifier (General)")]
[Tooltip("Multiplier applied on top of moveSpeed. 1.0 = default.")]
public float baseSpeedMultiplier = 1f;

[Header("Aura (General)")]
[Tooltip("Radius in world units. 0 disables aura.")]
public float auraRadius = 0f;
[Tooltip("Speed multiplier applied to affected same-era enemies.")]
public float auraSpeedMultiplier = 1.3f;""",
    )
    add_commit(doc, "feat(data): SALIN-54 extend EnemyDataSO with era, zigzag, and aura fields")

    # ---------- 2.4 PensionadoMover ----------
    add_h2(doc, "2.4  PensionadoMover component")
    add_para(
        doc,
        "Applies a horizontal sine offset on top of the base downward motion owned by Enemy. "
        "Auto-no-ops when zigzagAmplitude is 0, so the component is safe to attach but "
        "harmless on non-zigzag data.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Enemy/PensionadoMover.cs")
    add_code(
        doc,
        """using UnityEngine;

[RequireComponent(typeof(Enemy))]
public class PensionadoMover : MonoBehaviour
{
    private Enemy _enemy;
    private float _baseX;
    private float _spawnTime;

    private void OnEnable()
    {
        _enemy = GetComponent<Enemy>();
        _baseX = transform.position.x;
        _spawnTime = Time.time;
    }

    private void Update()
    {
        if (GameManager.Instance.CurrentState != GameState.Playing) return;

        var data = _enemy.Data;
        if (data == null || data.zigzagAmplitude <= 0f) return;

        float t = Time.time - _spawnTime;
        float offset = Mathf.Sin(t * Mathf.PI * 2f * data.zigzagFrequency)
                       * data.zigzagAmplitude;

        var pos = transform.position;
        pos.x = _baseX + offset;
        transform.position = pos;
    }
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "Downward motion is still driven by the existing Enemy mover. PensionadoMover only "
        "applies the horizontal offset. If Enemy.Update writes X directly every frame, "
        "PensionadoMover's write will be clobbered \u2014 in that case, add a flag like "
        "`_allowHorizontalOverride` on Enemy and have Enemy skip X writes when a mover has "
        "claimed control. One-line fix; test in-scene.",
    )
    add_commit(doc, "feat(enemy): SALIN-54 add PensionadoMover for zigzag variant")

    # ---------- 2.5 GeneralAura ----------
    add_h2(doc, "2.5  GeneralAura component")
    add_para(
        doc,
        "Ticks every 0.25s, iterates ActiveEnemyTracker.ActiveEnemies, applies the speed buff "
        "to American-era non-boss enemies within auraRadius. OnDisable cleanly removes the "
        "buff from anyone currently affected, so defeating the General removes the buff on "
        "the same frame.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Enemy/GeneralAura.cs")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;

[RequireComponent(typeof(Enemy))]
public class GeneralAura : MonoBehaviour
{
    private const float TICK = 0.25f;

    private Enemy _enemy;
    private readonly HashSet<Enemy> _affected = new HashSet<Enemy>();
    private readonly HashSet<Enemy> _stillAffectedBuffer = new HashSet<Enemy>();
    private float _nextTick;

    private void OnEnable()
    {
        _enemy = GetComponent<Enemy>();
        _nextTick = 0f;
    }

    private void OnDisable()
    {
        foreach (var e in _affected)
        {
            if (e != null) e.ClearSpeedBuff(this);
        }
        _affected.Clear();
    }

    private void Update()
    {
        if (GameManager.Instance.CurrentState != GameState.Playing) return;
        if (Time.time < _nextTick) return;
        _nextTick = Time.time + TICK;

        var data = _enemy.Data;
        if (data == null || data.auraRadius <= 0f) return;

        var active = ActiveEnemyTracker.Instance.GetActiveEnemiesSnapshot();
        var radiusSq = data.auraRadius * data.auraRadius;
        var myPos = transform.position;

        _stillAffectedBuffer.Clear();

        for (int i = 0; i < active.Count; i++)
        {
            var other = active[i];
            if (other == null || other == _enemy) continue;
            if (other.Data.era != EnemyDataSO.Era.American) continue;
            if (other.IsBoss) continue;

            float distSq = (other.transform.position - myPos).sqrMagnitude;
            if (distSq > radiusSq) continue;

            other.ApplySpeedBuff(this, data.auraSpeedMultiplier);
            _stillAffectedBuffer.Add(other);
        }

        // Remove buff from anything that was in our set last tick but is no longer.
        foreach (var prev in _affected)
        {
            if (prev != null && !_stillAffectedBuffer.Contains(prev))
                prev.ClearSpeedBuff(this);
        }

        _affected.Clear();
        foreach (var e in _stillAffectedBuffer) _affected.Add(e);
    }
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "GeneralAura uses GetActiveEnemiesSnapshot() \u2014 already available on the tracker "
        "courtesy of SALIN-89. It allocates a new list per tick, which is fine at a 0.25s "
        "cadence. Do NOT reuse FindAllWithCharacter's buffer here; the aura needs ALL active "
        "enemies, not a character-filtered view.",
    )
    add_commit(doc, "feat(enemy): SALIN-54 add GeneralAura proximity speed-buff component")

    # ---------- 2.6 Prefabs ----------
    add_h2(doc, "2.6  Author prefabs")
    add_para(doc, "Create two prefabs under Assets/Prefabs/Enemies/ based on the existing Enemy prefab:")
    add_bullets(
        doc,
        [
            "Enemy_Pensionado.prefab \u2014 duplicate Enemy.prefab, add PensionadoMover component.",
            "Enemy_General.prefab \u2014 duplicate Enemy.prefab, add GeneralAura component. Optional: add a disabled SpriteRenderer child scaled to 2\u00d7auraRadius as an in-editor aura gizmo (hidden in builds).",
        ],
    )
    add_commit(doc, "feat(enemy): SALIN-54 author Enemy_Pensionado and Enemy_General prefabs")

    # ---------- 2.7 SOs ----------
    add_h2(doc, "2.7  Author EnemyData SOs")
    add_para(doc, "In the Project window: Create \u2192 Salinlahi \u2192 EnemyData. Populate:")
    add_table(
        doc,
        ["Asset", "Key Fields"],
        [
            (
                "EnemyData_Pensionado.asset",
                "enemyID=Pensionado, era=American, isBoss=false, moveSpeed=0.9, maxHealth=1, "
                "assignedCharacter=(any Chapter-2 char), zigzagAmplitude=1.2, "
                "zigzagFrequency=2, baseSpeedMultiplier=1.0, auraRadius=0",
            ),
            (
                "EnemyData_General.asset",
                "enemyID=General, era=American, isBoss=false, moveSpeed=1.0, maxHealth=3, "
                "assignedCharacter=(any Chapter-2 char), zigzagAmplitude=0, "
                "baseSpeedMultiplier=0.7, auraRadius=3.5, auraSpeedMultiplier=1.3",
            ),
        ],
        col_widths=[Cm(5), Cm(12)],
    )
    add_commit(doc, "feat(data): SALIN-54 author EnemyData_Pensionado and EnemyData_General assets")

    # ---------- 2.8 EnemyPool mapping ----------
    add_h2(doc, "2.8  Register variants with EnemyPool")
    add_para(
        doc,
        "Open EnemyPool (or your enemy spawn registry) and add mappings: EnemyData_Pensionado "
        "\u2192 Enemy_Pensionado.prefab and EnemyData_General \u2192 Enemy_General.prefab. This is "
        "typically a list-of-pairs field on the pool SO or a Resources.Load convention \u2014 "
        "match your project's existing pattern.",
    )
    add_commit(doc, "chore(enemy): SALIN-54 register Pensionado and General variants with EnemyPool")

    # ---------- 2.9 Verification ----------
    add_h2(doc, "2.9  Verification")
    add_table(
        doc,
        ["Test", "Expected"],
        [
            (
                "Spawn 3 Pensionados in a test wave",
                "All three visibly zigzag horizontally while descending. All reach the shrine if undefeated.",
            ),
            (
                "Spawn 1 General plus 4 regular American-era Soldados in the General's radius",
                "Soldados visibly speed up (1.3\u00d7) while the General is alive.",
            ),
            (
                "Defeat the General while Soldados are still in radius",
                "Soldados revert to base speed on the same frame.",
            ),
            (
                "Spawn 1 General plus 2 Japanese-era enemies in radius",
                "Japanese enemies unchanged \u2014 aura is era-scoped.",
            ),
            (
                "Spawn 1 General plus 1 boss inside radius (if SALIN-68 is live)",
                "Boss unaffected by aura (IsBoss check).",
            ),
            (
                "Run 30 seconds with 5 Generals and 20 Soldados",
                "No GC spikes in Profiler's GC Alloc column (aura is allocation-free at tick).",
            ),
            (
                "Defeat a Pensionado",
                "Returns to pool and unregisters cleanly with single Unregister (fix from 2.1).",
            ),
        ],
        col_widths=[Cm(8), Cm(9)],
    )


# ===========================================================================
# SALIN-68  —  Boss Encounter System
# ===========================================================================


def task_salin_68(doc):
    add_h1(doc, "SALIN-68  Implement Boss Encounter System (BossConfigSO + Phase Mechanics)")
    add_para(
        doc,
        "Composition-over-inheritance boss system per the Boss Encounter spec "
        "(docs/superpowers/specs/2026-04-15-boss-encounter-system-design.md). One "
        "BossController with a phase state machine; small ability MonoBehaviours for "
        "boss-specific behavior. El Inquisidor ships as a full 3-phase encounter. "
        "Superintendent (Level 10) and Kadiliman (Level 15) ship as 1-phase placeholder stubs.",
    )
    add_para(
        doc,
        "This is the largest of the four tickets. Work the sub-steps in order \u2014 each unlocks "
        "the next. The internal order is: infrastructure \u2192 data types \u2192 runtime \u2192 UI \u2192 authored "
        "assets \u2192 tests.",
    )

    # ---------- 3.1 Assets/Tests scaffold ----------
    add_h2(doc, "3.1  Test assembly \u2014 already scaffolded by SALIN-89")
    add_callout(
        doc,
        "NOTE",
        "SALIN-89 already created Assets/Tests/Editor/Salinlahi.Tests.Editor.asmdef and "
        "shipped ProgressManagerTests, ActiveEnemyTrackerTests, ComboManagerTests under it. "
        "The tests use a SINGLE Editor asmdef (not an EditMode/PlayMode split). "
        "Write new edit-mode tests for step 3.18 inside Assets/Tests/Editor/Gameplay/.",
    )
    add_callout(
        doc,
        "WARNING",
        "The play-mode smoke test in step 3.19 needs a PlayMode asmdef that does not exist yet. "
        "Create ONLY Assets/Tests/PlayMode/Salinlahi.Tests.PlayMode.asmdef at that point \u2014 do "
        "not touch the existing Editor asmdef. Instructions below are kept for reference only; "
        "do not apply them verbatim.",
    )
    add_para(
        doc,
        "For reference, the PlayMode asmdef should look like this when you create it in 3.19:",
    )
    add_file_path(doc, "Assets/Tests/PlayMode/Salinlahi.Tests.PlayMode.asmdef")
    add_code(
        doc,
        """{
    "name": "Salinlahi.Tests.PlayMode",
    "rootNamespace": "",
    "references": [
        "UnityEngine.TestRunner",
        "UnityEditor.TestRunner",
        "Salinlahi.Runtime"
    ],
    "includePlatforms": [],
    "excludePlatforms": [],
    "allowUnsafeCode": false,
    "overrideReferences": true,
    "precompiledReferences": [
        "nunit.framework.dll"
    ],
    "autoReferenced": false,
    "defineConstraints": [ "UNITY_INCLUDE_TESTS" ],
    "optionalUnityReferences": [ "TestAssemblies" ]
}""",
    )
    add_para(
        doc,
        "No commit for step 3.1 \u2014 the Editor asmdef already exists. The PlayMode asmdef "
        "lands as part of step 3.19's commit.",
    )

    # ---------- 3.2 Singleton ----------
    add_h2(doc, "3.2  CombatResolver scene-lifecycle \u2014 already handled by SALIN-89")
    add_callout(
        doc,
        "NOTE",
        "SALIN-89 converted CombatResolver from Singleton<CombatResolver> to a plain "
        "MonoBehaviour in the Gameplay scene. It no longer has an Instance accessor, is "
        "not DontDestroyOnLoad, and Unity destroys it on scene unload \u2014 the stale-instance "
        "risk described by TICKETS-19/02 is gone. No Singleton.cs edit needed.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Combat/CombatResolver.cs")
    add_para(doc, "Current signature (for reference \u2014 do not edit here, step 3.11 edits it):")
    add_code(
        doc,
        """public class CombatResolver : MonoBehaviour
{
    private void OnEnable()  { EventBus.OnCharacterRecognized += HandleCharacterRecognized; }
    private void OnDisable() { EventBus.OnCharacterRecognized -= HandleCharacterRecognized; }
    private void HandleCharacterRecognized(string characterID) { ... }
}""",
    )
    add_callout(
        doc,
        "WARNING",
        "Because CombatResolver is not a Singleton anymore, step 3.11's boss routing must "
        "reach CurrentBoss via GameManager.Instance.CurrentBoss \u2014 NOT via any "
        "CombatResolver.Instance accessor (there isn't one). The code in 3.11 already uses "
        "GameManager.Instance.CurrentBoss; do not introduce a back-reference.",
    )
    add_para(doc, "No commit for step 3.2. Skip to 3.3.")

    # ---------- 3.3 EventBus boss events ----------
    add_h2(doc, "3.3  Declare boss lifecycle events on EventBus")
    add_para(
        doc,
        "Three events plus Raise helpers. Every downstream consumer (BGM, HUD, "
        "LevelFlowController, BossLabelIconRow) reads them. Declare them first so the rest "
        "of this task compiles.",
    )
    add_file_path(doc, "Assets/Scripts/Core/EventBus.cs")
    add_code(
        doc,
        """// Raised once when StartBoss is called. BossLabelIconRow, HUD, BGM swap all listen.
public static event Action<BossConfigSO> OnBossStarted;

// Raised per phase cleared, with the phase index that just ended.
public static event Action<int> OnBossPhaseCleared;

// Raised when the final phase clears. Outro plays; level complete follows outro.
public static event Action OnBossDefeated;

public static void RaiseBossStarted(BossConfigSO config)
{
    OnBossStarted?.Invoke(config);
}

public static void RaiseBossPhaseCleared(int phaseIndex)
{
    OnBossPhaseCleared?.Invoke(phaseIndex);
}

public static void RaiseBossDefeated()
{
    OnBossDefeated?.Invoke();
}""",
    )
    add_commit(doc, "feat(events): SALIN-68 declare boss lifecycle events on EventBus")

    # ---------- 3.4 GameManager.CurrentBoss ----------
    add_h2(doc, "3.4  Add GameManager.CurrentBoss")
    add_para(
        doc,
        "CombatResolver (step 3.11) and BossLabelIconRow (step 3.14) both read it. "
        "BossController sets it in StartBoss and clears it on outro completion.",
    )
    add_file_path(doc, "Assets/Scripts/Core/GameManager.cs")
    add_code(
        doc,
        """// Set by BossController.StartBoss; cleared on outro completion.
// Any code that needs to know whether a boss is active reads this.
public BossController CurrentBoss { get; set; }""",
    )
    add_commit(doc, "feat(core): SALIN-68 add CurrentBoss property to GameManager")

    # ---------- 3.5 ActiveEnemyTracker ActiveCount ----------
    add_h2(doc, "3.5  ActiveCount / IsClear \u2014 already on the tracker")
    add_callout(
        doc,
        "NOTE",
        "SALIN-89 already added both on Assets/Scripts/Gameplay/Enemy/ActiveEnemyTracker.cs. "
        "ActiveCount runs CleanupStaleEntries() under the hood, and IsClear \u21d4 ActiveCount == 0. "
        "Use them directly from BossController. No edit and no commit for step 3.5.",
    )
    add_code(
        doc,
        """// Already shipped:
public int ActiveCount { get { CleanupStaleEntries(); return _activeEnemies.Count; } }
public bool IsClear => ActiveCount == 0;""",
    )

    # ---------- 3.6 BaybayinCharacterSO.icon ----------
    add_h2(doc, "3.6  Add BaybayinCharacterSO.icon field")
    add_para(
        doc,
        "The boss label icon row (step 3.14) renders a 32\u00d732 icon per required character. "
        "Add the field and a fallback pattern that uses displaySprite if the icon is unassigned.",
    )
    add_file_path(doc, "Assets/Data/BaybayinCharacterSO.cs")
    add_code(
        doc,
        """[Tooltip("32x32 icon used by boss label row and UI toast. Falls back to displaySprite if null.")]
public Sprite icon;""",
    )
    add_callout(
        doc,
        "NOTE",
        "Unity deserializes existing .asset files cleanly after a new field is added; the "
        "icon will simply be null until you assign it. That is fine \u2014 BossLabelIconRow "
        "falls back to displaySprite. Author the 17 icons at your leisure; a stylized "
        "monochrome glyph is enough for MVP.",
    )
    add_commit(doc, "feat(data): SALIN-68 add icon sprite field to BaybayinCharacterSO")

    # ---------- 3.7 WaveSpawner.SpawnWave ----------
    add_h2(doc, "3.7  Extract WaveSpawner.SpawnWave")
    add_para(
        doc,
        "The spawn loop currently lives inside WaveManager.SpawnWaveRoutine. BossController "
        "spawns intermission waves by calling the spawner directly, which means the loop must "
        "be callable from outside WaveManager. Move it to WaveSpawner as a reusable coroutine.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Waves/WaveSpawner.cs")
    add_para(doc, "Add this public coroutine entry point:")
    add_code(
        doc,
        """public IEnumerator SpawnWave(WaveConfigSO wave)
{
    if (wave == null || wave.entries == null) yield break;

    foreach (var entry in wave.entries)
    {
        for (int i = 0; i < entry.count; i++)
        {
            SpawnEnemy(entry.enemyData);
            if (entry.spacing > 0f)
                yield return new WaitForSeconds(entry.spacing);
        }
        if (entry.trailingDelay > 0f)
            yield return new WaitForSeconds(entry.trailingDelay);
    }
}""",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Waves/WaveManager.cs")
    add_para(doc, "Then in WaveManager.SpawnWaveRoutine, replace the hand-rolled loop with a single yield:")
    add_code(
        doc,
        """private IEnumerator SpawnWaveRoutine(WaveConfigSO wave)
{
    EventBus.RaiseWaveStarted(wave);
    yield return _spawner.SpawnWave(wave);
    EventBus.RaiseWaveCompleted(wave);
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "If your WaveConfigSO uses different field names than entries / count / spacing / "
        "trailingDelay, adjust the loop body to match. The signature matters, not the "
        "exact inner fields.",
    )
    add_commit(doc, "refactor(waves): SALIN-68 extract spawn loop into WaveSpawner.SpawnWave")

    # ---------- 3.8 BossConfigSO ----------
    add_h2(doc, "3.8  BossConfigSO and BossPhase data types")
    add_file_path(doc, "Assets/Scripts/Data/BossConfigSO.cs")
    add_code(
        doc,
        """using System;
using System.Collections.Generic;
using UnityEngine;

public enum BossMovementPattern { Hover, Pace, Teleport }

[Serializable]
public class BossPhase
{
    [Tooltip("Every character must be drawn once to clear this phase \u2014 any order.")]
    public List<BaybayinCharacterSO> requiredCharacters;

    public BossMovementPattern movementPattern = BossMovementPattern.Hover;
    public float movementSpeed = 1f;

    [Tooltip("If set, this wave spawns after the phase clears. The next phase begins once the wave is fully defeated.")]
    public WaveConfigSO intermissionWave;

    [Tooltip("Grace period after intermission ends, before the next phase begins.")]
    public float postIntermissionDelay = 0.75f;
}

[CreateAssetMenu(menuName = "Salinlahi/BossConfig", fileName = "BossConfig_Name")]
public class BossConfigSO : ScriptableObject
{
    public string bossName;
    public Sprite bossSprite;
    public GameObject bossPrefab;
    public List<BossPhase> phases;
    public float introDuration = 1.25f;
    public float outroDuration = 1.5f;
}""",
    )
    add_commit(doc, "feat(boss): SALIN-68 add BossConfigSO and BossPhase data types")

    # ---------- 3.9 LevelConfigSO.bossConfig ----------
    add_h2(doc, "3.9  Add LevelConfigSO.bossConfig")
    add_file_path(doc, "Assets/Scripts/Data/LevelConfigSO.cs")
    add_code(
        doc,
        """[Header("Boss (optional)")]
[Tooltip("If set, this level is a boss encounter. WaveManager skips normal wave flow and hands off to BossController.")]
public BossConfigSO bossConfig;""",
    )
    add_commit(doc, "feat(data): SALIN-68 add bossConfig field to LevelConfigSO")

    # ---------- 3.10 BossController ----------
    add_h2(doc, "3.10  BossController state machine")
    add_para(
        doc,
        "Owns the phase state machine: Intro \u2192 PhaseActive \u2192 Intermission \u2192 \u2026 \u2192 Outro \u2192 "
        "Defeated. Raises EventBus boss events at each transition. TryHit consumes drawn "
        "characters and returns true when the boss claims the draw (right or wrong \u2014 duplicates "
        "and non-required chars are handled internally via OnDrawingFailed).",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Boss/BossController.cs")
    add_code(
        doc,
        """using System.Collections;
using System.Collections.Generic;
using UnityEngine;

public class BossController : MonoBehaviour
{
    public enum State { Idle, Intro, PhaseActive, Intermission, Outro, Defeated }

    public BossConfigSO Config { get; private set; }
    public State CurrentState { get; private set; } = State.Idle;
    public int PhaseIndex { get; private set; }
    public bool IsTargetable => CurrentState == State.PhaseActive;
    public BossPhase CurrentPhase =>
        (Config != null && PhaseIndex >= 0 && PhaseIndex < Config.phases.Count)
            ? Config.phases[PhaseIndex]
            : null;

    public event System.Action<BossPhase> OnPhaseStarted;
    public event System.Action<BossPhase> OnPhaseCleared;

    private readonly HashSet<string> _drawnThisPhase = new HashSet<string>();
    private bool _phaseComplete;

    public void StartBoss(BossConfigSO config)
    {
        Config = config;
        PhaseIndex = 0;
        GameManager.Instance.CurrentBoss = this;
        EventBus.RaiseBossStarted(config);
        StartCoroutine(RunEncounter());
    }

    // Returns true if the draw was consumed by the boss (right or wrong).
    // Returns false if the boss is untargetable or the char is not in this phase.
    public bool TryHit(BaybayinCharacterSO character)
    {
        if (!IsTargetable || CurrentPhase == null || character == null) return false;

        bool isRequired = CurrentPhase.requiredCharacters
            .Exists(c => c != null && c.characterID == character.characterID);
        if (!isRequired) return false;

        if (_drawnThisPhase.Contains(character.characterID))
        {
            EventBus.RaiseDrawingFailed();   // duplicate draw
            return true;
        }

        _drawnThisPhase.Add(character.characterID);

        if (_drawnThisPhase.Count >= CurrentPhase.requiredCharacters.Count)
            _phaseComplete = true;

        return true;
    }

    private IEnumerator RunEncounter()
    {
        // Intro
        CurrentState = State.Intro;
        yield return WaitPaused(Config.introDuration);

        while (PhaseIndex < Config.phases.Count)
        {
            CurrentState = State.PhaseActive;
            _drawnThisPhase.Clear();
            _phaseComplete = false;
            OnPhaseStarted?.Invoke(CurrentPhase);

            // Park until TryHit marks the phase complete.
            while (!_phaseComplete) yield return null;

            EventBus.RaiseBossPhaseCleared(PhaseIndex);
            OnPhaseCleared?.Invoke(CurrentPhase);

            bool isLast = PhaseIndex >= Config.phases.Count - 1;

            if (!isLast)
            {
                var wave = CurrentPhase.intermissionWave;
                if (wave != null)
                {
                    CurrentState = State.Intermission;
                    var spawner = FindObjectOfType<WaveSpawner>();
                    if (spawner != null) yield return spawner.SpawnWave(wave);

                    yield return new WaitUntil(() =>
                        ActiveEnemyTracker.Instance != null
                        && ActiveEnemyTracker.Instance.IsClear);

                    yield return WaitPaused(CurrentPhase.postIntermissionDelay);
                }

                PhaseIndex++;
                continue;
            }

            // Last phase clear \u2192 outro.
            EventBus.RaiseBossDefeated();
            CurrentState = State.Outro;
            yield return WaitPaused(Config.outroDuration);
            CurrentState = State.Defeated;
            GameManager.Instance.CurrentBoss = null;
            EventBus.RaiseLevelComplete();
            yield break;
        }
    }

    // Pausable wait. GameState.Paused freezes progression (AC-12).
    private IEnumerator WaitPaused(float seconds)
    {
        float t = 0f;
        while (t < seconds)
        {
            if (GameManager.Instance.CurrentState != GameState.Paused)
                t += Time.deltaTime;
            yield return null;
        }
    }
}""",
    )
    add_commit(doc, "feat(boss): SALIN-68 add BossController with phase state machine and pause gating")

    # ---------- 3.11 CombatResolver ----------
    add_h2(doc, "3.11  Add boss priority short-circuit in CombatResolver")
    add_para(
        doc,
        "Spec \u00a77: draws route to the boss first when one is present and targetable. If the "
        "boss's current phase requires the drawn character and has not yet consumed it, the "
        "hit goes to the boss and the resolver returns before AOE / closest-match logic runs.",
    )
    add_file_path(doc, "Assets/Scripts/Gameplay/Combat/CombatResolver.cs")
    add_para(doc, "Replace the OnCharacterRecognized handler (or equivalent) with this version:")
    add_code(
        doc,
        """private void OnCharacterRecognized(string characterID)
{
    var charSO = _registry.GetByID(characterID);
    if (charSO == null) return;

    // --- Boss priority (spec \u00a77) ---
    // Routes to the boss before AOE or closest-match. BossController.TryHit
    // handles duplicate and wrong-char internally (raises OnDrawingFailed).
    var boss = GameManager.Instance != null ? GameManager.Instance.CurrentBoss : null;
    if (boss != null && boss.IsTargetable)
    {
        if (boss.TryHit(charSO))
            return;
    }

    // Normal single-target path. AOEResolver (SALIN-40) runs on the same event
    // and no-ops when count < 3, so the two paths never double-hit (AOE AC-4).
    var target = ActiveEnemyTracker.Instance.FindClosestToBase(characterID);
    if (target != null)
        target.TakeDamage(1);
}""",
    )
    add_commit(doc, "feat(combat): SALIN-68 route boss draws before closest-match in CombatResolver")

    # ---------- 3.12 PhaseBasedMovement ----------
    add_h2(doc, "3.12  PhaseBasedMovement ability")
    add_file_path(doc, "Assets/Scripts/Gameplay/Boss/Abilities/PhaseBasedMovement.cs")
    add_code(
        doc,
        """using System.Collections;
using UnityEngine;

[RequireComponent(typeof(BossController))]
public class PhaseBasedMovement : MonoBehaviour
{
    [SerializeField] private float _paceDistance = 2.5f;
    [SerializeField] private Vector2[] _teleportPoints;

    private BossController _boss;
    private Coroutine _motion;
    private Vector3 _origin;

    private void Awake()
    {
        _boss = GetComponent<BossController>();
        _origin = transform.position;
    }

    private void OnEnable() { _boss.OnPhaseStarted += OnPhaseStarted; }
    private void OnDisable() { _boss.OnPhaseStarted -= OnPhaseStarted; }

    private void OnPhaseStarted(BossPhase phase)
    {
        if (_motion != null) StopCoroutine(_motion);
        switch (phase.movementPattern)
        {
            case BossMovementPattern.Hover:
                _motion = StartCoroutine(Hover(phase.movementSpeed)); break;
            case BossMovementPattern.Pace:
                _motion = StartCoroutine(Pace(phase.movementSpeed)); break;
            case BossMovementPattern.Teleport:
                _motion = StartCoroutine(Teleport(phase.movementSpeed)); break;
        }
    }

    private IEnumerator Hover(float speed)
    {
        float t = 0f;
        while (true)
        {
            if (GameManager.Instance.CurrentState == GameState.Paused) { yield return null; continue; }
            t += Time.deltaTime * speed;
            transform.position = _origin + new Vector3(0f, Mathf.Sin(t) * 0.25f, 0f);
            yield return null;
        }
    }

    private IEnumerator Pace(float speed)
    {
        float t = 0f;
        while (true)
        {
            if (GameManager.Instance.CurrentState == GameState.Paused) { yield return null; continue; }
            t += Time.deltaTime * speed;
            transform.position = _origin + new Vector3(Mathf.Sin(t) * _paceDistance, 0f, 0f);
            yield return null;
        }
    }

    private IEnumerator Teleport(float speed)
    {
        if (_teleportPoints == null || _teleportPoints.Length == 0) yield break;
        int i = 0;
        while (true)
        {
            if (GameManager.Instance.CurrentState == GameState.Paused) { yield return null; continue; }
            yield return new WaitForSeconds(1.2f / Mathf.Max(0.1f, speed));
            i = (i + 1) % _teleportPoints.Length;
            transform.position = _origin + (Vector3)_teleportPoints[i];
        }
    }
}""",
    )
    add_commit(doc, "feat(boss): SALIN-68 add PhaseBasedMovement ability component")

    # ---------- 3.13 SummonWaveOnPhaseStart ----------
    add_h2(doc, "3.13  SummonWaveOnPhaseStart ability")
    add_file_path(doc, "Assets/Scripts/Gameplay/Boss/Abilities/SummonWaveOnPhaseStart.cs")
    add_code(
        doc,
        """using UnityEngine;

[RequireComponent(typeof(BossController))]
public class SummonWaveOnPhaseStart : MonoBehaviour
{
    [Tooltip("Phase indices this ability should fire on.")]
    [SerializeField] private int[] _phaseIndices = { 0, 1 };
    [SerializeField] private WaveConfigSO _addsWave;

    private BossController _boss;

    private void Awake() { _boss = GetComponent<BossController>(); }

    private void OnEnable() { _boss.OnPhaseStarted += OnPhaseStarted; }
    private void OnDisable() { _boss.OnPhaseStarted -= OnPhaseStarted; }

    private void OnPhaseStarted(BossPhase phase)
    {
        if (System.Array.IndexOf(_phaseIndices, _boss.PhaseIndex) < 0) return;
        if (_addsWave == null) return;
        var spawner = FindObjectOfType<WaveSpawner>();
        if (spawner != null) StartCoroutine(spawner.SpawnWave(_addsWave));
    }
}""",
    )
    add_commit(doc, "feat(boss): SALIN-68 add SummonWaveOnPhaseStart ability component")

    # ---------- 3.14 BossLabelIconRow ----------
    add_h2(doc, "3.14  Boss label icon row UI")
    add_para(
        doc,
        "Renders the current phase's required characters as 32\u00d732 icons. Greys drawn ones "
        "to 40% alpha. Subscribes to OnBossStarted, OnCharacterRecognized, OnBossPhaseCleared, "
        "OnBossDefeated.",
    )
    add_file_path(doc, "Assets/Scripts/UI/BossLabelIconRow.cs")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;
using UnityEngine.UI;

public class BossLabelIconRow : MonoBehaviour
{
    [SerializeField] private Image _iconPrefab;
    [SerializeField] private RectTransform _row;
    [SerializeField] private int _maxPerRow = 6;
    [SerializeField] private float _drawnAlpha = 0.4f;

    private readonly List<Image> _icons = new List<Image>();
    private BossController _boss;

    private void OnEnable()
    {
        EventBus.OnBossStarted += OnBossStarted;
        EventBus.OnCharacterRecognized += OnRecognized;
        EventBus.OnBossPhaseCleared += _ => Clear();
        EventBus.OnBossDefeated += Clear;
    }

    private void OnDisable()
    {
        EventBus.OnBossStarted -= OnBossStarted;
        EventBus.OnCharacterRecognized -= OnRecognized;
    }

    private void OnBossStarted(BossConfigSO config)
    {
        _boss = GameManager.Instance.CurrentBoss;
        if (_boss != null) _boss.OnPhaseStarted += OnPhaseStarted;
    }

    private void OnPhaseStarted(BossPhase phase)
    {
        Clear();
        int count = Mathf.Min(phase.requiredCharacters.Count, _maxPerRow);
        for (int i = 0; i < count; i++)
        {
            var req = phase.requiredCharacters[i];
            var icon = Instantiate(_iconPrefab, _row);
            icon.sprite = req.icon != null ? req.icon : req.displaySprite;
            icon.color = Color.white;
            icon.name = $"Icon_{req.characterID}";
            _icons.Add(icon);
        }
    }

    private void OnRecognized(string characterID)
    {
        foreach (var icon in _icons)
        {
            if (icon == null) continue;
            if (icon.name == $"Icon_{characterID}")
            {
                var c = icon.color; c.a = _drawnAlpha; icon.color = c;
            }
        }
    }

    private void Clear()
    {
        foreach (var icon in _icons) if (icon != null) Destroy(icon.gameObject);
        _icons.Clear();
    }
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "AC-8 (from SALIN-40) requires AOE VFX to not obscure this row. Put the boss HUD "
        "canvas on a higher sort order than the gameplay FX canvas: Gameplay FX (order 5) < "
        "Boss HUD (order 10). Since AOEResolver already skips draws a boss would consume, "
        "the overlap is only an edge-case visual concern.",
    )
    add_commit(doc, "feat(ui): SALIN-68 add BossLabelIconRow for phase required characters")

    # ---------- 3.15 WaveManager.StartLevel branch ----------
    add_h2(doc, "3.15  WaveManager.StartLevel boss branch")
    add_file_path(doc, "Assets/Scripts/Gameplay/Waves/WaveManager.cs")
    add_para(doc, "Branch on the new bossConfig field before the normal wave loop begins:")
    add_code(
        doc,
        """public void StartLevel(LevelConfigSO level)
{
    _currentLevel = level;

    if (level.bossConfig != null)
    {
        var bossGO = Instantiate(level.bossConfig.bossPrefab);
        var controller = bossGO.GetComponent<BossController>();
        controller.StartBoss(level.bossConfig);
        return; // Skip normal wave flow \u2014 boss owns the level.
    }

    // Existing wave flow unchanged below this line.
    StartCoroutine(RunAllWavesRoutine());
}""",
    )
    add_commit(doc, "feat(waves): SALIN-68 branch StartLevel to BossController when bossConfig is set")

    # ---------- 3.16 El Inquisidor ----------
    add_h2(doc, "3.16  Author El Inquisidor")
    add_para(
        doc,
        "Create the boss prefab and BossConfigSO for the full 3-phase encounter (Level 5).",
    )
    add_h3(doc, "Prefab")
    add_para(
        doc,
        "Create Assets/Prefabs/Bosses/Boss_ElInquisidor.prefab. Attach BossController, "
        "PhaseBasedMovement (configure _paceDistance and _teleportPoints), and two "
        "SummonWaveOnPhaseStart components \u2014 one with phaseIndices=[0] and "
        "addsWave=Boss_L5_Intermission1, one with phaseIndices=[1] and "
        "addsWave=Boss_L5_Intermission2. Also add a SpriteRenderer with the boss art.",
    )
    add_h3(doc, "BossConfig")
    add_para(
        doc,
        "Create Assets/ScriptableObjects/Bosses/BossConfig_ElInquisidor.asset. Populate:",
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("bossName", "El Inquisidor"),
            ("bossSprite", "(assign boss art)"),
            ("bossPrefab", "Boss_ElInquisidor.prefab"),
            ("introDuration", "1.25"),
            ("outroDuration", "1.5"),
            (
                "phases[0]",
                "requiredCharacters=[BA], movementPattern=Hover, movementSpeed=1.0, "
                "intermissionWave=Boss_L5_Intermission1 (3 Soldados), "
                "postIntermissionDelay=0.75",
            ),
            (
                "phases[1]",
                "requiredCharacters=[KA], movementPattern=Pace, movementSpeed=1.5, "
                "intermissionWave=Boss_L5_Intermission2 (5 Soldados), "
                "postIntermissionDelay=0.75",
            ),
            (
                "phases[2]",
                "requiredCharacters=[GA], movementPattern=Teleport, movementSpeed=1.0, "
                "intermissionWave=(none), postIntermissionDelay=0",
            ),
        ],
        col_widths=[Cm(4.5), Cm(12.5)],
    )
    add_para(
        doc,
        "Also author the intermission waves (Boss_L5_Intermission1.asset, "
        "Boss_L5_Intermission2.asset) as WaveConfigSO assets under Assets/ScriptableObjects/Waves/. "
        "Set Level 5's LevelConfigSO.bossConfig to the El Inquisidor asset.",
    )
    add_commit(doc, "feat(boss): SALIN-68 author El Inquisidor prefab, config, and intermission waves")

    # ---------- 3.17 Stubs ----------
    add_h2(doc, "3.17  Author Superintendent and Kadiliman placeholder configs")
    add_para(doc, "Both are 1-phase stubs so Levels 10 and 15 load and are beatable. Reuse Boss_ElInquisidor.prefab as the prefab for now (palette swap later).")
    add_table(
        doc,
        ["Asset", "Configuration"],
        [
            (
                "BossConfig_Superintendent.asset",
                "bossName=Superintendent, introDuration=1.0, outroDuration=1.0, phases=[{requiredCharacters=[NA], movementPattern=Hover, movementSpeed=1.0, intermissionWave=null}]",
            ),
            (
                "BossConfig_Kadiliman.asset",
                "bossName=Kadiliman, introDuration=1.0, outroDuration=1.0, phases=[{requiredCharacters=[A], movementPattern=Hover, movementSpeed=1.0, intermissionWave=null}]",
            ),
        ],
        col_widths=[Cm(6), Cm(11)],
    )
    add_para(doc, "Wire Level 10 \u2192 BossConfig_Superintendent.asset and Level 15 \u2192 BossConfig_Kadiliman.asset.")
    add_commit(doc, "feat(boss): SALIN-68 author Superintendent and Kadiliman placeholder configs")

    # ---------- 3.18 Edit-mode tests ----------
    add_h2(doc, "3.18  Edit-mode unit tests")
    add_file_path(doc, "Assets/Tests/EditMode/BossControllerTests.cs")
    add_code(
        doc,
        """using System.Collections;
using NUnit.Framework;
using UnityEngine;
using UnityEngine.TestTools;

public class BossControllerTests
{
    private BossController _boss;
    private GameObject _host;

    private BaybayinCharacterSO MakeChar(string id)
    {
        var c = ScriptableObject.CreateInstance<BaybayinCharacterSO>();
        c.characterID = id;
        return c;
    }

    [SetUp]
    public void SetUp()
    {
        _host = new GameObject("BossHost");
        _boss = _host.AddComponent<BossController>();
    }

    [TearDown]
    public void TearDown() { Object.DestroyImmediate(_host); }

    [UnityTest]
    public IEnumerator SingleCharPhase_Advances_OnDraw()
    {
        var cfg = ScriptableObject.CreateInstance<BossConfigSO>();
        cfg.phases = new System.Collections.Generic.List<BossPhase>
        {
            new BossPhase { requiredCharacters = new() { MakeChar("BA") } },
            new BossPhase { requiredCharacters = new() { MakeChar("KA") } }
        };
        cfg.introDuration = 0f;
        _boss.StartBoss(cfg);
        yield return null; // let intro advance

        Assert.IsTrue(_boss.TryHit(MakeChar("BA")));
        yield return null;
        Assert.AreEqual(1, _boss.PhaseIndex);
    }

    [UnityTest]
    public IEnumerator MultiCharPhase_RequiresAllChars()
    {
        var cfg = ScriptableObject.CreateInstance<BossConfigSO>();
        cfg.phases = new() { new BossPhase { requiredCharacters = new() {
            MakeChar("BA"), MakeChar("KA") } } };
        cfg.introDuration = 0f;
        _boss.StartBoss(cfg);
        yield return null;

        Assert.IsTrue(_boss.TryHit(MakeChar("BA")));
        Assert.AreEqual(0, _boss.PhaseIndex); // still in phase 0
        Assert.IsTrue(_boss.TryHit(MakeChar("KA")));
        yield return null;
        // After last phase, outro runs; state transitions to Outro then Defeated.
    }

    [UnityTest]
    public IEnumerator DuplicateDraw_RaisesFailedAndDoesNotAdvance()
    {
        bool failed = false;
        System.Action h = () => failed = true;
        EventBus.OnDrawingFailed += h;

        var cfg = ScriptableObject.CreateInstance<BossConfigSO>();
        cfg.phases = new() { new BossPhase { requiredCharacters = new() {
            MakeChar("BA"), MakeChar("KA") } } };
        cfg.introDuration = 0f;
        _boss.StartBoss(cfg);
        yield return null;

        _boss.TryHit(MakeChar("BA"));
        _boss.TryHit(MakeChar("BA")); // duplicate
        Assert.IsTrue(failed);

        EventBus.OnDrawingFailed -= h;
    }

    [UnityTest]
    public IEnumerator IntroIsUntargetable()
    {
        var cfg = ScriptableObject.CreateInstance<BossConfigSO>();
        cfg.phases = new() { new BossPhase { requiredCharacters = new() { MakeChar("BA") } } };
        cfg.introDuration = 5f;
        _boss.StartBoss(cfg);
        Assert.IsFalse(_boss.IsTargetable);
        yield return null;
        Assert.IsFalse(_boss.IsTargetable);
    }

    [UnityTest]
    public IEnumerator LastPhase_FiresBossDefeated()
    {
        bool defeated = false;
        System.Action h = () => defeated = true;
        EventBus.OnBossDefeated += h;

        var cfg = ScriptableObject.CreateInstance<BossConfigSO>();
        cfg.phases = new() { new BossPhase { requiredCharacters = new() { MakeChar("BA") } } };
        cfg.introDuration = 0f;
        cfg.outroDuration = 0f;
        _boss.StartBoss(cfg);
        yield return null;
        _boss.TryHit(MakeChar("BA"));
        yield return null;
        yield return null;

        Assert.IsTrue(defeated);
        EventBus.OnBossDefeated -= h;
    }
}""",
    )
    add_commit(doc, "test(boss): SALIN-68 edit-mode tests for BossController phase machine")

    # ---------- 3.19 Play-mode test ----------
    add_h2(doc, "3.19  Play-mode smoke test")
    add_file_path(doc, "Assets/Tests/PlayMode/ElInquisidorSmokeTest.cs")
    add_code(
        doc,
        """using System.Collections;
using NUnit.Framework;
using UnityEngine;
using UnityEngine.SceneManagement;
using UnityEngine.TestTools;

public class ElInquisidorSmokeTest
{
    [UnityTest]
    public IEnumerator ElInquisidor_CompletesEndToEnd()
    {
        bool levelComplete = false;
        System.Action h = () => levelComplete = true;
        EventBus.OnLevelComplete += h;

        yield return SceneManager.LoadSceneAsync("Gameplay");
        // Level loader should auto-select Level 5 for this test.
        // If it does not, drive GameManager.StartLevel(levelFiveConfig) here.

        yield return new WaitUntil(() => GameManager.Instance.CurrentBoss != null);
        var boss = GameManager.Instance.CurrentBoss;

        yield return new WaitUntil(() => boss.IsTargetable);
        EventBus.RaiseCharacterRecognized("BA");
        yield return new WaitUntil(() =>
            boss.CurrentState == BossController.State.Intermission
            || boss.PhaseIndex > 0);
        yield return new WaitUntil(() => boss.IsTargetable);
        EventBus.RaiseCharacterRecognized("KA");
        yield return new WaitUntil(() => boss.PhaseIndex > 1 || boss.IsTargetable == false);
        yield return new WaitUntil(() => boss.IsTargetable);
        EventBus.RaiseCharacterRecognized("GA");

        yield return new WaitUntil(() => levelComplete);
        Assert.IsTrue(levelComplete);

        EventBus.OnLevelComplete -= h;
    }
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "Intermission waves need to be short (or mocked) so the test runs under the play-mode "
        "timeout. If the real intermission spawns take too long, author test-only WaveConfigSO "
        "assets with count=1 and assign them in a test fixture.",
    )
    add_commit(doc, "test(boss): SALIN-68 play-mode smoke test for El Inquisidor end-to-end")

    # ---------- 3.20 Verification ----------
    add_h2(doc, "3.20  Verification")
    add_table(
        doc,
        ["Test", "Expected"],
        [
            ("Enter Level 5", "Boss spawns, intro plays, first phase activates. Icon row shows BA."),
            ("Draw BA during phase 0", "Phase advances. Intermission 3-Soldado wave spawns."),
            ("Defeat all Soldados", "Phase 1 activates. Icon row switches to KA. Boss uses Pace movement."),
            ("Pause mid-phase", "Boss motion freezes. WaitPaused timers stop accumulating."),
            ("Draw correct then same char again", "Duplicate draw raises OnDrawingFailed. Phase does not advance."),
            ("Complete all 3 phases", "OnBossDefeated fires. Outro plays for 1.5s. OnLevelComplete fires."),
            ("Run EditMode tests", "All 5 tests green."),
            ("Run PlayMode smoke test", "Green under 60s."),
            ("Enter Level 10", "Superintendent 1-phase encounter loads and is beatable."),
            ("Enter Level 15", "Kadiliman 1-phase encounter loads and is beatable."),
        ],
        col_widths=[Cm(7.5), Cm(9.5)],
    )


# ===========================================================================
# SALIN-69  —  Tracing Dojo
# ===========================================================================


def task_salin_69(doc):
    add_h1(doc, "SALIN-69  Implement Tracing Dojo Scene and Practice Mode")
    add_para(
        doc,
        "A dedicated practice scene reachable from the Main Menu. Players draw characters "
        "with no enemies, no hearts, no wave timer. Selecting a character shows a ghost "
        "stroke overlay. Each attempt returns a confidence % and pass/fail verdict; on pass, "
        "the pronunciation clip plays. Practice attempts must NOT touch the SALIN-35 research "
        "log.",
    )
    add_callout(
        doc,
        "NOTE",
        "Partial scaffolding already exists on the feature/SALIN-69-tracing-dojo branch: the "
        "TracingDojo scene file, the four dojo scripts, and the Main Menu button handler are "
        "all on disk. This section splits into (a) verify-only steps for what is already "
        "shipped, (b) a Unity-Editor authoring pass for the TracingDojo scene hierarchy, and "
        "(c) three code fixes that make the existing scripts compile against the real APIs "
        "(BaybayinCharacterSO has no `icon`, RecognitionManager has no `LastResult`, and the "
        "populator needs a CharacterRegistrySO because no list-of-characters asset exists).",
    )

    # ---------- 4.1 RecognitionLogger (already shipped) ----------
    add_h2(doc, "4.1  Verify RecognitionLogger.LoggingEnabled (already shipped)")
    add_para(
        doc,
        "Shipped in commit 4f34295 on this branch. No code changes and no commit in this "
        "step \u2014 verify only.",
    )
    add_file_path(doc, "Assets/Scripts/Analytics/RecognitionLogger.cs")
    add_bullets(
        doc,
        [
            "(a) Project window \u2192 locate Assets/Scripts/Analytics/RecognitionLogger.cs \u2192 double-click to open.",
            "(b) Confirm line 9 declares: public static bool LoggingEnabled { get; set; } = true;",
            "(c) Confirm line 32 (inside LogAttempt) has the guard: if (!LoggingEnabled) return;",
            "(d) Close the file. No edit, no commit.",
        ],
    )
    add_callout(
        doc,
        "NOTE",
        "Default value is true so gameplay and SALIN-35 logging keep working identically. The "
        "dojo controller (\u00a74.3) is the only consumer that flips it \u2014 false on OnEnable, back "
        "to true on OnDisable.",
    )

    # ---------- 4.2 Scene hierarchy ----------
    add_h2(doc, "4.2  Author the TracingDojo scene hierarchy")
    add_para(
        doc,
        "The scene file Assets/_Scenes/TracingDojo.unity already exists and is already listed "
        "in Build Settings (index 6 in ProjectSettings/EditorBuildSettings.asset). This step "
        "builds its hierarchy click-by-click. Menu paths cite Unity 6 LTS (6000.3.9f1).",
    )
    add_callout(
        doc,
        "WARNING",
        "The scenes folder in this project is Assets/_Scenes/ (with a leading underscore). "
        "Earlier drafts of this guide said Assets/Scenes/ \u2014 that folder does not exist.",
    )

    add_h3(doc, "4.2.1  Open the scene and clean defaults")
    add_bullets(
        doc,
        [
            "(a) File \u2192 Open Scene \u2192 navigate to Assets/_Scenes/TracingDojo.unity \u2192 Open.",
            "(b) If the Hierarchy window contains a Directional Light GameObject, right-click it \u2192 Delete. The dojo is a 2D screen-space scene and does not need scene lighting.",
            "(c) Leave Main Camera in place. Select it \u2192 Inspector \u2192 Camera component \u2192 set Projection = Orthographic if not already. Background color: dark grey (40, 40, 50, 255).",
        ],
    )

    add_h3(doc, "4.2.2  Create the Canvas")
    add_bullets(
        doc,
        [
            "(a) Hierarchy window \u2192 right-click the empty area \u2192 UI \u2192 Canvas. Unity creates two GameObjects at the scene root: `Canvas` (with Canvas + Canvas Scaler + Graphic Raycaster components) and `EventSystem` (handles UI input).",
            "(b) Select Canvas. Inspector \u2192 Canvas component \u2192 Render Mode dropdown (currently `Screen Space - Overlay`) \u2192 choose `Screen Space - Camera`. A new `Render Camera` field appears.",
            "(c) Drag `Main Camera` from the Hierarchy into the `Render Camera` slot. Plane Distance stays at 100.",
            "(d) Inspector \u2192 Canvas Scaler component \u2192 click `UI Scale Mode` dropdown (currently `Constant Pixel Size`) \u2192 choose `Scale With Screen Size`. The rows below change.",
            "(e) Reference Resolution row \u2192 set X = 1080, Y = 1920 (portrait phone target).",
            "(f) Screen Match Mode \u2192 choose `Match Width Or Height`. Drag the `Match` slider to 0.5 (or type 0.5 in the numeric field). Leave Reference Pixels Per Unit at 100.",
        ],
    )

    add_h3(doc, "4.2.3  Create CharacterList (left column scroll view)")
    add_bullets(
        doc,
        [
            "(a) Right-click Canvas in the Hierarchy \u2192 UI \u2192 Scroll View. Unity adds `Scroll View` with children Viewport \u2192 Content plus two scrollbars. Rename `Scroll View` to `CharacterList` (F2).",
            "(b) Select CharacterList. RectTransform \u2192 click the anchor preset icon (top-left square of RectTransform). Hold Alt+Shift and click the stretch preset, then manually set Anchor Min = (0, 0), Anchor Max = (0.35, 1), Left/Top/Right/Bottom = 0. This pins it to the left 35% of the screen.",
            "(c) Inside CharacterList, select `Scrollbar Horizontal` \u2192 right-click \u2192 Delete. We only need vertical scrolling.",
            "(d) Select the remaining `Scrollbar Vertical`. RectTransform anchor preset \u2192 stretch-right. Leave Width = 20.",
            "(e) On CharacterList, select the Scroll Rect component \u2192 uncheck `Horizontal`. Drag the remaining `Scrollbar Vertical` into the `Vertical Scrollbar` field if it is not already bound.",
            "(f) Expand CharacterList \u2192 Viewport \u2192 Content. Select Content.",
            "(g) Content \u2192 Add Component \u2192 `Vertical Layout Group` (arranges children top-to-bottom and computes preferred height). Set Spacing = 8, Child Alignment = Upper Center. Under `Control Child Size`, check Width, leave Height unchecked. Under `Child Force Expand`, check Width, leave Height unchecked.",
            "(h) Content \u2192 Add Component \u2192 `Content Size Fitter` (resizes Content to fit children so the Scroll Rect knows its scrollable extent). Set Vertical Fit = Preferred Size. Leave Horizontal Fit = Unconstrained.",
        ],
    )

    add_h3(doc, "4.2.4  Create DrawingPanel (right side) and GhostStrokeLayer")
    add_bullets(
        doc,
        [
            "(a) Right-click Canvas \u2192 UI \u2192 Image. Rename the new child to `DrawingPanel`.",
            "(b) Select DrawingPanel \u2192 RectTransform \u2192 Anchor Min = (0.35, 0), Anchor Max = (1, 1), Left/Top/Right/Bottom = 0. This fills the right 65% of the screen.",
            "(c) Inspector \u2192 Image component \u2192 Color swatch \u2192 set (20, 20, 30, 255) \u2014 dark blue-black. Leave `Raycast Target` checked so drawing input receives hits.",
            "(d) Right-click DrawingPanel \u2192 Create Empty. Rename to `GhostStrokeLayer`. This is the RectTransform the GhostStrokeRenderer paints into.",
            "(e) Select GhostStrokeLayer \u2192 RectTransform \u2192 stretch to parent (Anchor Min = (0, 0), Anchor Max = (1, 1), Left/Top/Right/Bottom = 0).",
        ],
    )

    add_h3(doc, "4.2.5  Create FeedbackPanel (verdict + confidence)")
    add_bullets(
        doc,
        [
            "(a) Right-click DrawingPanel \u2192 Create Empty. Rename to `FeedbackPanel`.",
            "(b) RectTransform \u2192 anchor preset: bottom-stretch (hold Alt, click the bottom-stretch preset). Height = 200, Pos Y = 24, Left/Right = 24.",
            "(c) Add Component \u2192 `Vertical Layout Group`. Child Alignment = Upper Center, Spacing = 8, Control Child Size Width/Height = both checked, Child Force Expand Width = checked.",
            "(d) Right-click FeedbackPanel \u2192 UI \u2192 Text - TextMeshPro. If a `TMP Essentials` import prompt appears, click `Import TMP Essentials`. Rename child to `VerdictLabel`.",
            "(e) Select VerdictLabel. TMP_Text \u2192 clear the placeholder text. Font Size = 72. Alignment = Center + Middle. Auto Size = off. Vertex Color = white (255, 255, 255, 255).",
            "(f) Right-click FeedbackPanel \u2192 UI \u2192 Text - TextMeshPro. Rename to `ConfidenceLabel`. Font Size = 48, Alignment = Center + Middle, Vertex Color = (200, 200, 200, 255). Clear placeholder text.",
        ],
    )

    add_h3(doc, "4.2.6  Create BackButton")
    add_bullets(
        doc,
        [
            "(a) Right-click Canvas \u2192 UI \u2192 Button - TextMeshPro. Unity creates a `Button` GameObject plus a `Text (TMP)` child. Rename the Button to `BackButton`.",
            "(b) Select BackButton \u2192 RectTransform \u2192 anchor preset top-left. Width = 200, Height = 80, Pos X = 60, Pos Y = -60.",
            "(c) Expand BackButton \u2192 select the child `Text (TMP)` \u2192 change Text to `\u2190 Back`. Font Size = 40.",
            "(d) Button wiring is covered in \u00a74.2.8 after the Dojo GameObject is set up.",
        ],
    )

    add_h3(doc, "4.2.7  Create the Dojo root and attach controllers")
    add_para(
        doc,
        "This is a scene-root empty GameObject that hosts the three dojo MonoBehaviours plus a "
        "small navigator helper for the Back button. It lets all the scripts sit on one object "
        "and share serialized references.",
    )
    add_bullets(
        doc,
        [
            "(a) Hierarchy \u2192 right-click empty area \u2192 Create Empty. Name it `Dojo`. Transform Reset (click the \u22ee icon on Transform \u2192 Reset) \u2192 position (0, 0, 0).",
            "(b) With Dojo selected, Inspector \u2192 Add Component \u2192 type `TracingDojoController` \u2192 Enter. This is the rewritten controller from \u00a74.3.",
            "(c) Add Component \u2192 type `GhostStrokeRenderer` \u2192 Enter.",
            "(d) Add Component \u2192 type `CharacterListPopulator` \u2192 Enter.",
            "(e) Add Component \u2192 type `DojoNavigator` \u2192 Enter. (You will create this script in \u00a74.2.8.)",
        ],
    )

    add_h3(doc, "4.2.8  Create DojoNavigator and wire the Back button")
    add_para(
        doc,
        "SceneLoader is a DontDestroyOnLoad singleton bootstrapped from Bootstrap. Adding a "
        "tiny scene-local navigator lets the Back button's OnClick target a method on a "
        "GameObject that exists in the TracingDojo scene, which is simpler and more robust "
        "than trying to reference a DDOL object from the scene.",
    )
    add_file_path(doc, "Assets/Scripts/UI/TracingDojo/DojoNavigator.cs")
    add_code(
        doc,
        """using UnityEngine;

public class DojoNavigator : MonoBehaviour
{
    public void GoToMainMenu()
    {
        SceneLoader.Instance.LoadScene(\"MainMenu\");
    }
}""",
    )
    add_bullets(
        doc,
        [
            "(a) Project window \u2192 right-click Assets/Scripts/UI/TracingDojo/ \u2192 Create \u2192 C# Script \u2192 name `DojoNavigator`. Paste the body above. Save.",
            "(b) Return to Unity and wait for the status bar to finish `Compiling...`.",
            "(c) Select Dojo in the Hierarchy \u2192 the DojoNavigator component added in \u00a74.2.7(e) is now resolved (no yellow warning).",
            "(d) Select BackButton in the Hierarchy. Inspector \u2192 Button component \u2192 `On Click ()` list \u2192 press the `+` button. A new row appears with an empty GameObject slot and a function dropdown reading `No Function`.",
            "(e) Drag the `Dojo` GameObject from the Hierarchy into the empty GameObject slot.",
            "(f) Click the function dropdown \u2192 `DojoNavigator` \u2192 `GoToMainMenu()`.",
        ],
    )

    add_h3(doc, "4.2.9  Add DrawingInput")
    add_bullets(
        doc,
        [
            "(a) Project window \u2192 search `DrawingInput` or locate the drawing input prefab used in Assets/_Scenes/Gameplay.unity. Drag it from the Project window into the Hierarchy so it becomes a child of Canvas (or whichever parent Gameplay uses \u2014 mirror the Gameplay scene exactly to avoid drift).",
            "(b) If the prefab has any gameplay-only dependencies (e.g. HeartManager, WaveManager references), they will log null warnings in the dojo. Those are expected; do NOT wire them \u2014 the dojo deliberately omits those systems.",
        ],
    )

    add_h3(doc, "4.2.10  Save the scene")
    add_bullets(
        doc,
        [
            "(a) File \u2192 Save (Ctrl+S). Confirm the scene tab title drops the modified-asterisk.",
            "(b) Commit after \u00a74.3 lands so the scene wiring and the script fixes ship together.",
        ],
    )

    # ---------- 4.3 Fix compile errors ----------
    add_h2(doc, "4.3  Fix compile errors in the on-disk dojo scripts")
    add_para(
        doc,
        "The four dojo scripts committed to this branch reference three symbols that do not "
        "exist in the current codebase: BaybayinCharacterSO.icon (field), "
        "BaybayinCharacterSO.introducedInLevelID (field), and RecognitionManager.LastResult "
        "(property). This step creates a new CharacterRegistrySO asset, swaps the missing "
        "field references for existing ones, and rewires TracingDojoController to use the "
        "OnRecognitionResolved event (which carries the RecognitionResult directly).",
    )

    add_h3(doc, "4.3.1  Create CharacterRegistrySO (new script)")
    add_para(
        doc,
        "Both CharacterListPopulator and TracingDojoController expect a CharacterRegistrySO "
        "field. No such asset exists yet \u2014 add it as a simple list-of-characters "
        "ScriptableObject. It is genuinely useful beyond the dojo (any future UI that needs "
        "to iterate all characters can reuse it).",
    )
    add_file_path(doc, "Assets/Scripts/Data/CharacterRegistrySO.cs")
    add_code(
        doc,
        """using System.Collections.Generic;
using UnityEngine;

[CreateAssetMenu(
    fileName = \"CharacterRegistry\",
    menuName = \"Salinlahi/Character Registry\")]
public class CharacterRegistrySO : ScriptableObject
{
    public List<BaybayinCharacterSO> All = new List<BaybayinCharacterSO>();
}""",
    )
    add_bullets(
        doc,
        [
            "(a) Project window \u2192 right-click Assets/Scripts/Data/ \u2192 Create \u2192 C# Script \u2192 name `CharacterRegistrySO`. Paste the body above. Save.",
            "(b) Return to Unity and wait for `Compiling...` to clear.",
            "(c) Project window \u2192 navigate to Assets/ScriptableObjects/Characters/ \u2192 right-click \u2192 Create \u2192 Salinlahi \u2192 Character Registry. Name the asset `CharacterRegistry_Default`.",
            "(d) Select CharacterRegistry_Default \u2192 Inspector \u2192 `All` list \u2192 click the size field, type the total number of Char_* assets (17 at the time of writing) \u2192 Enter. The list expands with empty element slots.",
            "(e) Select all Char_*.asset files in Assets/ScriptableObjects/Characters/ (Ctrl+click each, or shift-click a range) and drag them into the All list. Unity assigns them in drop order \u2014 reorder to alphabetical if you want deterministic list ordering in the dojo.",
            "(f) Save the project (Ctrl+S, or File \u2192 Save Project).",
        ],
    )

    add_h3(doc, "4.3.2  Fix CharacterListRow.cs (icon \u2192 displaySprite)")
    add_file_path(doc, "Assets/Scripts/UI/TracingDojo/CharacterListRow.cs")
    add_para(
        doc,
        "BaybayinCharacterSO does not declare an `icon` field. The existing on-disk row file "
        "references `character.icon` twice on lines 14\u201316. Replace the three-line conditional "
        "with a single assignment to `character.displaySprite`.",
    )
    add_code(
        doc,
        """using System;
using TMPro;
using UnityEngine;
using UnityEngine.UI;

public class CharacterListRow : MonoBehaviour
{
    [SerializeField] private Image _iconImage;
    [SerializeField] private TMP_Text _label;
    [SerializeField] private Button _button;

    public void Bind(BaybayinCharacterSO character, Action<BaybayinCharacterSO> onSelect)
    {
        _iconImage.sprite = character.displaySprite;
        _label.text = character.characterID;
        _button.onClick.RemoveAllListeners();
        _button.onClick.AddListener(() => onSelect(character));
    }
}""",
    )

    add_h3(doc, "4.3.3  Fix CharacterListPopulator.cs (remove introducedInLevelID gate)")
    add_file_path(doc, "Assets/Scripts/UI/TracingDojo/CharacterListPopulator.cs")
    add_para(
        doc,
        "BaybayinCharacterSO does not declare `introducedInLevelID`. During development all "
        "characters in the registry are practice-eligible \u2014 iterate the list and skip nothing. "
        "When per-character unlock gating is designed, add it to ProgressManager (or a "
        "dedicated CharacterProgressSO) and reintroduce the filter here.",
    )
    add_code(
        doc,
        """using UnityEngine;

public class CharacterListPopulator : MonoBehaviour
{
    [SerializeField] private CharacterRegistrySO _registry;
    [SerializeField] private CharacterListRow _rowPrefab;
    [SerializeField] private Transform _content;
    [SerializeField] private TracingDojoController _controller;

    private void Start()
    {
        foreach (Transform child in _content) Destroy(child.gameObject);

        foreach (var character in _registry.All)
        {
            var row = Instantiate(_rowPrefab, _content);
            row.Bind(character, _controller.SelectCharacter);
        }
    }
}""",
    )

    add_h3(doc, "4.3.4  Rewrite TracingDojoController.cs (use OnRecognitionResolved)")
    add_file_path(doc, "Assets/Scripts/UI/TracingDojo/TracingDojoController.cs")
    add_para(
        doc,
        "RecognitionManager does not expose a `LastResult` property. EventBus does, however, "
        "raise `OnRecognitionResolved(RecognitionResult result, bool passedThreshold, float "
        "threshold)` for every attempt \u2014 both passes and fails. Subscribing to that single "
        "event gives the controller everything it needs in one place: the full result, the "
        "pass/fail verdict, and the threshold. This also eliminates the two separate "
        "OnCharacterRecognized + OnDrawingFailed handlers.",
    )
    add_code(
        doc,
        """using TMPro;
using UnityEngine;

public class TracingDojoController : MonoBehaviour
{
    [SerializeField] private GhostStrokeRenderer _ghost;
    [SerializeField] private TMP_Text _verdictLabel;
    [SerializeField] private TMP_Text _confidenceLabel;
    [SerializeField] private CharacterRegistrySO _registry;

    private BaybayinCharacterSO _selected;

    private void OnEnable()
    {
        EventBus.OnRecognitionResolved += OnResolved;
        RecognitionLogger.LoggingEnabled = false; // practice attempts are not logged
    }

    private void OnDisable()
    {
        EventBus.OnRecognitionResolved -= OnResolved;
        RecognitionLogger.LoggingEnabled = true;
    }

    public void SelectCharacter(BaybayinCharacterSO character)
    {
        _selected = character;
        _ghost.Render(character);
        _verdictLabel.text = string.Empty;
        _confidenceLabel.text = string.Empty;
    }

    private void OnResolved(RecognitionResult result, bool passedThreshold, float threshold)
    {
        RenderFeedback(result.characterID, result.score, passedThreshold);

        if (passedThreshold
            && _selected != null
            && result.characterID == _selected.characterID
            && _selected.pronunciationClip != null)
        {
            AudioManager.Instance.PlaySFX(_selected.pronunciationClip);
        }
    }

    private void RenderFeedback(string characterID, float score, bool pass)
    {
        _confidenceLabel.text = $\"{score * 100f:F0}%\";
        _verdictLabel.text = pass
            ? $\"\u2713  {characterID}\"
            : $\"\u2717  Try again ({characterID}?)\";
        _verdictLabel.color = pass
            ? new Color(0.20f, 0.55f, 0.25f)
            : new Color(0.70f, 0.20f, 0.20f);
    }
}""",
    )
    add_callout(
        doc,
        "NOTE",
        "The RecognitionConfigSO field was dropped \u2014 the dojo no longer needs it because the "
        "threshold it used to read from that asset is now carried on every "
        "OnRecognitionResolved payload as the third argument. Removing a dead serialized field "
        "keeps Inspector wiring honest.",
    )
    add_commit(
        doc,
        "fix(dojo): SALIN-69 wire dojo scripts to existing APIs and CharacterRegistrySO",
    )
    add_para(
        doc,
        "Scene authoring from \u00a74.2 rides on this same commit since the scene references the "
        "rewritten scripts and the new CharacterRegistrySO asset.",
    )
    add_commit(
        doc,
        "feat(dojo): SALIN-69 author TracingDojo scene hierarchy and component wiring",
    )

    # ---------- 4.4 GhostStrokeRenderer (verify + dot prefab) ----------
    add_h2(doc, "4.4  Verify GhostStrokeRenderer and author GhostDot prefab")
    add_para(
        doc,
        "GhostStrokeRenderer.cs is already on disk and compiles against current APIs \u2014 no "
        "code changes. What it needs is (a) a UI dot prefab bound to _dotPrefab, and (b) "
        "template files on disk matching the `Templates/{characterID}_template` convention.",
    )
    add_file_path(doc, "Assets/Scripts/UI/TracingDojo/GhostStrokeRenderer.cs")
    add_bullets(
        doc,
        [
            "(a) Open the file and scan. Confirm the Render method builds `Templates/{character.characterID}_template` via Resources.Load<TextAsset>. Confirm Parse splits on newline, splits on comma, parses two floats as (x, y) in [0, 1].",
            "(b) Project window \u2192 Assets/Resources/Templates/. Confirm that BA_template.txt, KA_template.txt, and GA_template.txt exist (canonical templates). Variant files such as A_template_01.txt exist but will not be loaded by the renderer \u2014 that is intentional; the dojo ghost overlay uses only the canonical template per character.",
            "(c) If a character in CharacterRegistry_Default lacks a canonical `{ID}_template.txt` file, the ghost will not render for that character and the renderer will log `no template for {ID}`. Either author the canonical template or temporarily remove that character from the registry.",
        ],
    )

    add_h3(doc, "4.4.1  Author the GhostDot prefab")
    add_bullets(
        doc,
        [
            "(a) Hierarchy (with any scene open) \u2192 right-click \u2192 UI \u2192 Image. Rename it `GhostDot`.",
            "(b) Select GhostDot \u2192 RectTransform \u2192 Pivot = (0.5, 0.5), Anchor Min = (0.5, 0.5), Anchor Max = (0.5, 0.5), Width = 8, Height = 8.",
            "(c) Image component \u2192 Color = (255, 255, 255, 255) white (alpha is re-applied at runtime by GhostStrokeRenderer per the `_ghostAlpha` slider).",
            "(d) Image component \u2192 uncheck `Raycast Target` (the dots must not block drawing input).",
            "(e) Drag GhostDot from the Hierarchy into Assets/Prefabs/UI/ (create the folder if it does not exist via right-click \u2192 Create \u2192 Folder). Unity creates GhostDot.prefab and converts the scene GameObject to a prefab instance.",
            "(f) Delete the now-prefab-instance GhostDot from the scene \u2014 the prefab is all we need.",
        ],
    )

    add_h3(doc, "4.4.2  Wire GhostStrokeRenderer serialized fields")
    add_bullets(
        doc,
        [
            "(a) Open Assets/_Scenes/TracingDojo.unity. Select the Dojo GameObject.",
            "(b) Inspector \u2192 GhostStrokeRenderer component \u2192 `Canvas Area` slot \u2192 drag Canvas/DrawingPanel/GhostStrokeLayer from the Hierarchy into it.",
            "(c) `Dot Prefab` slot \u2192 drag Assets/Prefabs/UI/GhostDot.prefab from the Project window into it.",
            "(d) `Ghost Alpha` slider \u2192 leave at 0.35.",
            "(e) Save scene (Ctrl+S).",
        ],
    )
    add_commit(
        doc,
        "feat(dojo): SALIN-69 author GhostDot prefab for ghost stroke overlay",
    )

    # ---------- 4.5 (folded into 4.3) ----------
    add_h2(doc, "4.5  TracingDojoController (covered in \u00a74.3.4)")
    add_para(
        doc,
        "The rewritten controller body and serialized-field wiring live in \u00a74.3.4. No separate "
        "commit \u2014 the controller fix ships in the same commit as the other script swaps.",
    )

    # ---------- 4.6 Main Menu button (already wired) ----------
    add_h2(doc, "4.6  Verify the Main Menu Practice button (already wired)")
    add_para(
        doc,
        "Assets/Scripts/UI/MainMenuUI.cs already exposes `OnTracingDojoPressed()` that calls "
        "`SceneLoader.Instance.LoadScene(\\\"TracingDojo\\\")`. This step verifies a button in "
        "the Main Menu scene is bound to that method, or wires one if missing.",
    )
    add_file_path(doc, "Assets/Scripts/UI/MainMenuUI.cs")
    add_bullets(
        doc,
        [
            "(a) Open Assets/_Scenes/MainMenu.unity.",
            "(b) Hierarchy \u2192 locate the button group under Canvas (likely named `Buttons` or `MainMenuButtons` \u2014 confirm by inspection; this scene is owned elsewhere in the project).",
            "(c) Look for an existing child button labeled `Practice` or `TracingDojo` or `Dojo`. If one already exists, select it \u2192 Inspector \u2192 Button component \u2192 On Click () \u2192 confirm the target function is `MainMenuUI \u2192 OnTracingDojoPressed()`. If it is, skip to (g).",
            "(d) If no such button exists: right-click the button group GameObject \u2192 UI \u2192 Button - TextMeshPro. Rename to `PracticeButton`. Reorder in the Vertical Layout Group to sit below `PlayButton` (drag in the Hierarchy).",
            "(e) Match the new button's RectTransform (Width, Height) to the existing siblings. If the parent has a Vertical Layout Group with Control Child Size enabled, the layout group sets the size automatically.",
            "(f) Select PracticeButton \u2192 Inspector \u2192 Button \u2192 On Click () \u2192 `+` \u2192 drag the GameObject that has MainMenuUI attached (likely named `MainMenu` or `MainMenuUI`) into the object slot \u2192 function dropdown \u2192 `MainMenuUI \u2192 OnTracingDojoPressed()`.",
            "(g) Select the button's `Text (TMP)` child \u2192 set text to `Practice`.",
            "(h) Save scene. If (c) confirmed an existing wire, there is no commit for this step. Otherwise use the commit below.",
        ],
    )
    add_callout(
        doc,
        "NOTE",
        "Earlier drafts of this guide named a non-existent class `MainMenuButtons` and method "
        "`OnPracticeClicked`. The real entry point is `MainMenuUI.OnTracingDojoPressed` at "
        "Assets/Scripts/UI/MainMenuUI.cs line 62. It is already callable \u2014 only the UI button "
        "binding may need authoring.",
    )
    add_commit(
        doc,
        "chore(menu): SALIN-69 wire Practice button to MainMenuUI.OnTracingDojoPressed (if not already bound)",
    )

    # ---------- 4.7 Character list row prefab ----------
    add_h2(doc, "4.7  Author CharacterListRow prefab and wire the populator")
    add_para(
        doc,
        "The populator's _rowPrefab slot expects a prefab with `CharacterListRow` attached and "
        "serialized references to an Image (icon), a TMP_Text (label), and a Button. This "
        "step authors that prefab and wires the populator's four serialized fields.",
    )

    add_h3(doc, "4.7.1  Build the row prefab in a scratch scene")
    add_bullets(
        doc,
        [
            "(a) Open Assets/_Scenes/TracingDojo.unity (we will author the row in this scene and extract it to a prefab).",
            "(b) Hierarchy \u2192 right-click Canvas \u2192 UI \u2192 Button - TextMeshPro. Rename to `CharacterListRow`. (Using a Button as the root gives us the Button + Image components for free.)",
            "(c) RectTransform \u2192 stretch horizontally inside its parent layout group: Anchor Min = (0, 1), Anchor Max = (1, 1), Pivot Y = 1, Width = auto (controlled by Vertical Layout Group), Height = 120.",
            "(d) The existing child `Text (TMP)` becomes the label. Select it, rename to `Label`. RectTransform: stretch to parent, Left = 140, Right = 16, Top = 0, Bottom = 0. Font Size = 40, Alignment = Middle Left, Text = blank (it is populated at runtime).",
            "(e) Right-click CharacterListRow \u2192 UI \u2192 Image. Rename to `Icon`. RectTransform: anchor preset left-middle. Width = 96, Height = 96, Pos X = 60, Pos Y = 0. Image Source Image can be left None \u2014 the populator assigns character.displaySprite at runtime.",
            "(f) Select CharacterListRow (the root) \u2192 Add Component \u2192 `CharacterListRow` (the MonoBehaviour we fixed in \u00a74.3.2).",
            "(g) On CharacterListRow component: drag Icon into `_iconImage`, drag Label into `_label`, drag the CharacterListRow GameObject itself (which has the Button component) into `_button`.",
            "(h) Drag CharacterListRow from the Hierarchy into Assets/Prefabs/UI/. Unity creates CharacterListRow.prefab.",
            "(i) Delete the now-prefab-instance CharacterListRow from the scene.",
        ],
    )

    add_h3(doc, "4.7.2  Wire CharacterListPopulator serialized fields")
    add_bullets(
        doc,
        [
            "(a) Select Dojo in the Hierarchy.",
            "(b) Inspector \u2192 CharacterListPopulator component \u2192 `_registry` \u2192 drag Assets/ScriptableObjects/Characters/CharacterRegistry_Default.asset into the slot.",
            "(c) `_rowPrefab` \u2192 drag Assets/Prefabs/UI/CharacterListRow.prefab.",
            "(d) `_content` \u2192 drag Canvas/CharacterList/Viewport/Content (from the Hierarchy) into the slot.",
            "(e) `_controller` \u2192 drag Dojo itself (since TracingDojoController lives on the same GameObject) into the slot.",
        ],
    )

    add_h3(doc, "4.7.3  Wire TracingDojoController serialized fields")
    add_bullets(
        doc,
        [
            "(a) Select Dojo. Inspector \u2192 TracingDojoController component.",
            "(b) `_ghost` \u2192 drag Dojo (its own GhostStrokeRenderer sibling) into the slot.",
            "(c) `_verdictLabel` \u2192 drag Canvas/DrawingPanel/FeedbackPanel/VerdictLabel.",
            "(d) `_confidenceLabel` \u2192 drag Canvas/DrawingPanel/FeedbackPanel/ConfidenceLabel.",
            "(e) `_registry` \u2192 drag CharacterRegistry_Default.asset.",
            "(f) Save scene (Ctrl+S).",
        ],
    )
    add_commit(
        doc,
        "feat(dojo): SALIN-69 author CharacterListRow prefab and wire populator",
    )

    # ---------- 4.8 Verification ----------
    add_h2(doc, "4.8  Verification")
    add_para(
        doc,
        "Before running: confirm the Unity Console shows zero red errors after saving the "
        "three fixed scripts and the new CharacterRegistrySO. If any red appears, re-read "
        "\u00a74.3 before continuing.",
    )
    add_table(
        doc,
        ["Test", "Expected"],
        [
            (
                "From Bootstrap, let the Main Menu load. Tap the Practice button.",
                "TracingDojo scene loads. CharacterList on the left is populated with rows (one per Char_* in CharacterRegistry_Default).",
            ),
            (
                "Tap the BA row.",
                "DrawingPanel shows a ~35% alpha scatter of dots over the canvas area in the shape of BA. VerdictLabel and ConfidenceLabel are blank.",
            ),
            (
                "Draw BA reasonably carefully inside DrawingPanel.",
                "ConfidenceLabel shows a percentage at or above RecognitionConfig.minimumConfidence. VerdictLabel shows green `\u2713  BA`. BA pronunciation plays once via AudioManager.PlaySFX.",
            ),
            (
                "Draw a scribble.",
                "ConfidenceLabel shows a low %. VerdictLabel shows red `\u2717  Try again (<topID>?)`. No audio.",
            ),
            (
                "Confirm the scene has no HUD elements.",
                "No HeartDisplay, ComboDisplay, WaveDisplay, or FocusModeIndicator present in the scene.",
            ),
            (
                "Measure RecognitionLogger row count before and after a 10-draw practice session. CSV path (Editor): %AppData%/../LocalLow/{company}/{product}/recognition_log.csv.",
                "Row count unchanged. Practice attempts did not log.",
            ),
            (
                "Tap Back. From Main Menu, Play through one Gameplay level and draw any character.",
                "CSV row count increases. SALIN-35 research logging is unaffected by dojo activity.",
            ),
        ],
        col_widths=[Cm(8), Cm(9)],
    )


# ===========================================================================
# Back matter
# ===========================================================================


def back_matter(doc):
    add_h1(doc, "Final Checklist")
    add_para(doc, "Walk through every item before closing the sprint. Grouped by Jira ticket.")

    add_h3(doc, "SALIN-40  AOE Burst Mechanic")
    add_checklist(
        doc,
        [
            "EventBus.OnAOETriggered declared with Raise helper.",
            "ActiveEnemyTracker allocation-free FindAllWithCharacter + CleanupStaleEntries verified (shipped in SALIN-89).",
            "ActiveEnemyTracker.GetActiveEnemiesSnapshot() available for GeneralAura-style iteration (shipped in SALIN-89).",
            "AOEResolver subscribes to OnCharacterRecognized in OnEnable, unsubscribes in OnDisable.",
            "Count >= 3 triggers TakeDamage(maxHealth) on every matching non-boss enemy in one pass.",
            "Count < 3 no-ops; CombatResolver handles the single-target path.",
            "OnAOETriggered raised exactly once per AOE, after damage lands.",
            "AOE never affects non-matching characters or bosses.",
            "AOEResolver GameObject placed in Gameplay scene (not Bootstrap).",
            "Defeated enemies return to pool and unregister cleanly.",
        ],
    )

    add_h3(doc, "SALIN-54  Pensionado + General")
    add_checklist(
        doc,
        [
            "Enemy.Defeat no longer double-unregisters.",
            "Enemy has EffectiveSpeed, ApplySpeedBuff, ClearSpeedBuff, IsBoss, MaxHealth.",
            "Enemy Update uses EffectiveSpeed (not Data.moveSpeed directly).",
            "EnemyDataSO has era, isBoss, zigzagAmplitude/Frequency, baseSpeedMultiplier, auraRadius, auraSpeedMultiplier.",
            "Pensionado produces visible non-linear zigzag and reaches the shrine if undefeated.",
            "General moves at 0.7\u00d7 base speed.",
            "American-era enemies in General's aura move at 1.3\u00d7.",
            "Aura buff removed from all affected enemies on General defeat.",
            "Aura never affects non-American enemies or bosses.",
            "Enemy_Pensionado.prefab and Enemy_General.prefab exist.",
            "EnemyData_Pensionado.asset and EnemyData_General.asset exist.",
            "EnemyPool maps both new variants.",
        ],
    )

    add_h3(doc, "SALIN-68  Boss Encounter System")
    add_checklist(
        doc,
        [
            "Assets/Tests EditMode and PlayMode asmdefs present.",
            "Singleton<T> clears _instance on destroy; CombatResolver is scene-local.",
            "EventBus declares OnBossStarted, OnBossPhaseCleared, OnBossDefeated with Raise helpers.",
            "GameManager.CurrentBoss property exists.",
            "ActiveEnemyTracker.ActiveCount and IsClear expose tracker state.",
            "BaybayinCharacterSO has an icon : Sprite field.",
            "WaveSpawner.SpawnWave(WaveConfigSO) exists; WaveManager delegates to it.",
            "BossConfigSO and BossPhase exist with documented fields.",
            "LevelConfigSO.bossConfig exists; WaveManager.StartLevel branches to boss.",
            "BossController state machine: Intro \u2192 PhaseActive \u2192 Intermission \u2192 \u2026 \u2192 Outro \u2192 Defeated.",
            "IsTargetable = false during Intro, Intermission, Outro.",
            "Phase clears when all requiredCharacters drawn once (any order).",
            "CombatResolver routes to BossController.TryHit before AOE / closest-match.",
            "Wrong or duplicate draws raise OnDrawingFailed.",
            "Last phase fires OnBossDefeated; after outro, OnLevelComplete.",
            "PhaseBasedMovement and SummonWaveOnPhaseStart abilities work from BossConfig data.",
            "BossLabelIconRow renders current phase's required characters (max 6 per row, drawn greyed).",
            "GameState.Paused halts all boss + ability coroutines.",
            "El Inquisidor full 3-phase encounter (BA, KA, GA; Hover \u2192 Pace \u2192 Teleport; 2 intermissions).",
            "Superintendent and Kadiliman 1-phase stubs load and are beatable.",
            "EditMode tests cover 5 cases.",
            "PlayMode smoke test completes El Inquisidor end-to-end.",
        ],
    )

    add_h3(doc, "SALIN-69  Tracing Dojo")
    add_checklist(
        doc,
        [
            "RecognitionLogger.LoggingEnabled flag gates LogAttempt (verified from commit 4f34295).",
            "Assets/_Scenes/TracingDojo.unity is in Build Settings (verified, already shipped).",
            "Main Menu Practice button is bound to MainMenuUI.OnTracingDojoPressed.",
            "CharacterRegistrySO script and CharacterRegistry_Default asset exist.",
            "Dojo GameObject hosts TracingDojoController, GhostStrokeRenderer, CharacterListPopulator, and DojoNavigator.",
            "GhostDot.prefab and CharacterListRow.prefab exist under Assets/Prefabs/UI/.",
            "TracingDojoController subscribes to EventBus.OnRecognitionResolved (not to OnCharacterRecognized / OnDrawingFailed).",
            "Selecting a character renders the template point cloud as a faded ghost overlay at ~35% alpha.",
            "Each attempt shows confidence % and pass/fail verdict based on the threshold carried by OnRecognitionResolved.",
            "Pass plays BaybayinCharacterSO.pronunciationClip via AudioManager.PlaySFX.",
            "Scene contains no HeartDisplay, ComboDisplay, WaveDisplay, or FocusModeIndicator.",
            "RecognitionLogger row count is unchanged before/after a practice session.",
            "Gameplay logging still works after returning from dojo (flag restored to true on OnDisable).",
        ],
    )

    add_h1(doc, "Files Created or Modified")
    add_table(
        doc,
        ["File", "Action", "Owning Ticket"],
        [
            ("Assets/Scripts/Core/EventBus.cs", "Modified", "SALIN-40 + SALIN-68"),
            ("Assets/Scripts/Gameplay/Enemy/ActiveEnemyTracker.cs", "Modified", "SALIN-40 + SALIN-68"),
            ("Assets/Scripts/Gameplay/Combat/AOEResolver.cs", "New", "SALIN-40"),
            ("Assets/Scenes/Gameplay.unity", "Modified", "SALIN-40"),
            ("Assets/Scripts/Gameplay/Enemy/Enemy.cs", "Modified", "SALIN-54"),
            ("Assets/Scripts/Data/EnemyDataSO.cs", "Modified", "SALIN-54"),
            ("Assets/Scripts/Gameplay/Enemy/PensionadoMover.cs", "New", "SALIN-54"),
            ("Assets/Scripts/Gameplay/Enemy/GeneralAura.cs", "New", "SALIN-54"),
            ("Assets/Prefabs/Enemies/Enemy_Pensionado.prefab", "New", "SALIN-54"),
            ("Assets/Prefabs/Enemies/Enemy_General.prefab", "New", "SALIN-54"),
            ("Assets/ScriptableObjects/Enemies/EnemyData_Pensionado.asset", "New", "SALIN-54"),
            ("Assets/ScriptableObjects/Enemies/EnemyData_General.asset", "New", "SALIN-54"),
            ("Assets/Tests/EditMode/Salinlahi.EditMode.asmdef", "New", "SALIN-68"),
            ("Assets/Tests/PlayMode/Salinlahi.PlayMode.asmdef", "New", "SALIN-68"),
            ("Assets/Scripts/Utilities/Singleton.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Core/GameManager.cs", "Modified", "SALIN-68"),
            ("Assets/Data/BaybayinCharacterSO.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Waves/WaveSpawner.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Waves/WaveManager.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Combat/CombatResolver.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Data/BossConfigSO.cs", "New", "SALIN-68"),
            ("Assets/Scripts/Data/LevelConfigSO.cs", "Modified", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Boss/BossController.cs", "New", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Boss/Abilities/PhaseBasedMovement.cs", "New", "SALIN-68"),
            ("Assets/Scripts/Gameplay/Boss/Abilities/SummonWaveOnPhaseStart.cs", "New", "SALIN-68"),
            ("Assets/Scripts/UI/BossLabelIconRow.cs", "New", "SALIN-68"),
            ("Assets/Prefabs/Bosses/Boss_ElInquisidor.prefab", "New", "SALIN-68"),
            ("Assets/ScriptableObjects/Bosses/BossConfig_ElInquisidor.asset", "New", "SALIN-68"),
            ("Assets/ScriptableObjects/Bosses/BossConfig_Superintendent.asset", "New", "SALIN-68"),
            ("Assets/ScriptableObjects/Bosses/BossConfig_Kadiliman.asset", "New", "SALIN-68"),
            ("Assets/Tests/EditMode/BossControllerTests.cs", "New", "SALIN-68"),
            ("Assets/Tests/PlayMode/ElInquisidorSmokeTest.cs", "New", "SALIN-68"),
            ("Assets/Scripts/Analytics/RecognitionLogger.cs", "Already shipped (commit 4f34295)", "SALIN-69"),
            ("Assets/_Scenes/TracingDojo.unity", "Already shipped (hierarchy authored in \u00a74.2)", "SALIN-69"),
            ("ProjectSettings/EditorBuildSettings.asset", "Already shipped (scene at index 6)", "SALIN-69"),
            ("Assets/Scripts/UI/TracingDojo/CharacterListPopulator.cs", "Modified (\u00a74.3.3)", "SALIN-69"),
            ("Assets/Scripts/UI/TracingDojo/CharacterListRow.cs", "Modified (\u00a74.3.2)", "SALIN-69"),
            ("Assets/Scripts/UI/TracingDojo/GhostStrokeRenderer.cs", "Already shipped (verify only)", "SALIN-69"),
            ("Assets/Scripts/UI/TracingDojo/TracingDojoController.cs", "Modified (\u00a74.3.4)", "SALIN-69"),
            ("Assets/Scripts/UI/TracingDojo/DojoNavigator.cs", "New (\u00a74.2.8)", "SALIN-69"),
            ("Assets/Scripts/Data/CharacterRegistrySO.cs", "New (\u00a74.3.1)", "SALIN-69"),
            ("Assets/ScriptableObjects/Characters/CharacterRegistry_Default.asset", "New (\u00a74.3.1)", "SALIN-69"),
            ("Assets/Prefabs/UI/GhostDot.prefab", "New (\u00a74.4.1)", "SALIN-69"),
            ("Assets/Prefabs/UI/CharacterListRow.prefab", "New (\u00a74.7.1)", "SALIN-69"),
            ("Assets/Scripts/UI/MainMenuUI.cs", "Already shipped (button wire verified in \u00a74.6)", "SALIN-69"),
            ("Assets/_Scenes/MainMenu.unity", "Modified only if Practice button is missing (\u00a74.6)", "SALIN-69"),
        ],
        col_widths=[Cm(9.5), Cm(3), Cm(4.5)],
    )

    add_h1(doc, "Commit Order")
    add_para(
        doc,
        "Each ticket is self-contained. You can land them in any order. Within a ticket, keep "
        "the listed sub-step order \u2014 later sub-steps in the same ticket consume code added by "
        "earlier ones.",
    )
    add_h3(doc, "SALIN-40 commits")
    add_bullets(
        doc,
        [
            "1.1  feat(events): SALIN-40 add OnAOETriggered event to EventBus",
            "1.2  (skipped \u2014 already shipped in SALIN-89; verify only)",
            "1.3  feat(combat): SALIN-40 add AOEResolver for 3+ same-character mass defeat",
            "1.4  chore(scene): SALIN-40 add AOEResolver GameObject to Gameplay scene",
            "1.5  feat(ui): SALIN-40 show mass-clear HUD badge on OnAOETriggered (optional)",
            "1.6  test(combat): SALIN-40 verify AOE threshold, exclusions, and boss priority",
        ],
    )
    add_h3(doc, "SALIN-54 commits")
    add_bullets(
        doc,
        [
            "2.1  fix(enemy): SALIN-54 remove redundant ActiveEnemyTracker.Unregister in Enemy.Defeat",
            "2.2  feat(enemy): SALIN-54 add stackable speed buff API, EffectiveSpeed, and IsBoss helper to Enemy",
            "2.3  feat(data): SALIN-54 extend EnemyDataSO with era, zigzag, and aura fields",
            "2.4  feat(enemy): SALIN-54 add PensionadoMover for zigzag variant",
            "2.5  feat(enemy): SALIN-54 add GeneralAura proximity speed-buff component",
            "2.6  feat(enemy): SALIN-54 author Enemy_Pensionado and Enemy_General prefabs",
            "2.7  feat(data): SALIN-54 author EnemyData_Pensionado and EnemyData_General assets",
            "2.8  chore(enemy): SALIN-54 register Pensionado and General variants with EnemyPool",
        ],
    )
    add_h3(doc, "SALIN-68 commits")
    add_bullets(
        doc,
        [
            "3.1  (skipped \u2014 Editor asmdef shipped in SALIN-89; PlayMode asmdef rolls into 3.19)",
            "3.2  (skipped \u2014 CombatResolver already scene-local after SALIN-89)",
            "3.3  feat(events): SALIN-68 declare boss lifecycle events on EventBus",
            "3.4  feat(core): SALIN-68 add CurrentBoss property to GameManager",
            "3.5  (skipped \u2014 ActiveCount / IsClear shipped in SALIN-89)",
            "3.6  feat(data): SALIN-68 add icon sprite field to BaybayinCharacterSO",
            "3.7  refactor(waves): SALIN-68 extract spawn loop into WaveSpawner.SpawnWave",
            "3.8  feat(boss): SALIN-68 add BossConfigSO and BossPhase data types",
            "3.9  feat(data): SALIN-68 add bossConfig field to LevelConfigSO",
            "3.10 feat(boss): SALIN-68 add BossController with phase state machine and pause gating",
            "3.11 feat(combat): SALIN-68 route boss draws before closest-match in CombatResolver",
            "3.12 feat(boss): SALIN-68 add PhaseBasedMovement ability component",
            "3.13 feat(boss): SALIN-68 add SummonWaveOnPhaseStart ability component",
            "3.14 feat(ui): SALIN-68 add BossLabelIconRow for phase required characters",
            "3.15 feat(waves): SALIN-68 branch StartLevel to BossController when bossConfig is set",
            "3.16 feat(boss): SALIN-68 author El Inquisidor prefab, config, and intermission waves",
            "3.17 feat(boss): SALIN-68 author Superintendent and Kadiliman placeholder configs",
            "3.18 test(boss): SALIN-68 edit-mode tests for BossController phase machine",
            "3.19 test(boss): SALIN-68 play-mode smoke test for El Inquisidor end-to-end",
        ],
    )
    add_h3(doc, "SALIN-69 commits")
    add_bullets(
        doc,
        [
            "4.1  (skipped \u2014 LoggingEnabled shipped in 4f34295; verify only)",
            "4.3  fix(dojo): SALIN-69 wire dojo scripts to existing APIs and CharacterRegistrySO",
            "4.3  feat(dojo): SALIN-69 author TracingDojo scene hierarchy and component wiring",
            "4.4  feat(dojo): SALIN-69 author GhostDot prefab for ghost stroke overlay",
            "4.6  chore(menu): SALIN-69 wire Practice button to MainMenuUI.OnTracingDojoPressed (only if not already bound)",
            "4.7  feat(dojo): SALIN-69 author CharacterListRow prefab and wire populator",
        ],
    )
    add_callout(
        doc,
        "TIP",
        "Tag after all four tickets merge: git tag -a v0.3.0 -m \"Sprint 2: AOE, Enemies, Boss, Dojo\".",
    )


if __name__ == "__main__":
    build()
